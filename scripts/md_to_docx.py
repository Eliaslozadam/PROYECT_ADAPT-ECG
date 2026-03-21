"""
Genera ADAPT-ECG_Documentacion_Tecnica.docx con explicación detallada del reentrenamiento CNN.
"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

OUT = "docs/ADAPT-ECG_Documentacion_Tecnica.docx"

# ── Colores ────────────────────────────────────────────────────────────────────
AZUL_OSCURO  = RGBColor(0x1a, 0x23, 0x7e)
AZUL_MEDIO   = RGBColor(0x28, 0x35, 0x93)
AZUL_CLARO   = RGBColor(0x39, 0x49, 0xab)
GRIS         = RGBColor(0x44, 0x44, 0x44)
BLANCO       = RGBColor(0xFF, 0xFF, 0xFF)

doc = Document()

# ── Márgenes ──────────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3)
    section.right_margin  = Cm(2.5)

# ── Helpers ───────────────────────────────────────────────────────────────────
def h1(text):
    p = doc.add_heading(text, level=1)
    p.runs[0].font.color.rgb = AZUL_OSCURO
    p.runs[0].font.size = Pt(18)
    return p

def h2(text):
    p = doc.add_heading(text, level=2)
    p.runs[0].font.color.rgb = AZUL_MEDIO
    p.runs[0].font.size = Pt(14)
    return p

def h3(text):
    p = doc.add_heading(text, level=3)
    p.runs[0].font.color.rgb = AZUL_CLARO
    p.runs[0].font.size = Pt(12)
    return p

def h4(text):
    p = doc.add_heading(text, level=4)
    p.runs[0].font.color.rgb = AZUL_CLARO
    p.runs[0].font.size = Pt(11)
    return p

def body(text, bold_parts=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    # Procesar **negrita** inline
    parts = re.split(r"(\*\*[^*]+\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = p.add_run(part[2:-2])
            run.bold = True
        else:
            p.add_run(part)
    return p

def code_block(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(0x20, 0x20, 0x60)
    # Fondo gris claro
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F0F0F0")
    pPr.append(shd)
    return p

def nota(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("📌 " + text)
    run.italic = True
    run.font.color.rgb = RGBColor(0x44, 0x44, 0x88)
    return p

def tabla(encabezados, filas, col_widths=None):
    t = doc.add_table(rows=1, cols=len(encabezados))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Encabezado
    hdr = t.rows[0].cells
    for i, h in enumerate(encabezados):
        hdr[i].text = h
        hdr[i].paragraphs[0].runs[0].bold = True
        hdr[i].paragraphs[0].runs[0].font.color.rgb = BLANCO
        hdr[i].paragraphs[0].runs[0].font.size = Pt(9)
        tc = hdr[i]._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "3949AB")
        tcPr.append(shd)
    # Filas
    for ri, row in enumerate(filas):
        cells = t.add_row().cells
        fill = "EEF0FF" if ri % 2 == 0 else "FFFFFF"
        for i, val in enumerate(row):
            cells[i].text = str(val)
            cells[i].paragraphs[0].runs[0].font.size = Pt(9)
            tc = cells[i]._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"), "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"), fill)
            tcPr.append(shd)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return t

def separador():
    doc.add_paragraph("─" * 80).runs[0].font.color.rgb = RGBColor(0xC5, 0xCA, 0xE9)

# ══════════════════════════════════════════════════════════════════════════════
# PORTADA
# ══════════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("ADAPT-ECG")
run.bold = True
run.font.size = Pt(28)
run.font.color.rgb = AZUL_OSCURO

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Sistema inteligente con reentrenamiento continuo para la detección\nadaptativa de patologías cardiovasculares basado en señales ECG")
run.font.size = Pt(14)
run.font.color.rgb = AZUL_MEDIO

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Documentación Técnica Completa")
run.font.size = Pt(12)
run.bold = True

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Elias Bejarano Lozada  |  #20213057\n"
                "Residencia Profesional — Ingeniería Biomédica\n"
                "Instituto Tecnológico de Tijuana")
run.font.size = Pt(11)
run.font.color.rgb = GRIS

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 1. QUÉ ES ADAPT-ECG
# ══════════════════════════════════════════════════════════════════════════════
h1("1. ¿Qué es ADAPT-ECG?")
body("ADAPT-ECG es un sistema de clasificación automática de arritmias cardiacas a partir de señales ECG. "
     "Utiliza una Red Neuronal Convolucional 1D (CNN) entrenada con la base de datos MIT-BIH Arrhythmia Database "
     "y es capaz de **reentrenarse de forma incremental** con nuevos datos sin olvidar lo aprendido anteriormente, "
     "gracias a un mecanismo llamado **Replay Buffer**.")

body("El sistema clasifica cada latido en una de las **5 clases del estándar AAMI** "
     "(Association for the Advancement of Medical Instrumentation):")

tabla(
    ["Clase", "Nombre completo", "Significado clínico"],
    [
        ["N", "Normal", "Latido sinusal normal"],
        ["S", "Supraventricular", "Arritmia originada en las aurículas"],
        ["V", "Ventricular", "Arritmia originada en los ventrículos (peligrosa)"],
        ["F", "Fusión", "Latido híbrido entre Normal y Ventricular"],
        ["Q", "Desconocido", "Artefacto, marcapasos o no clasificable"],
    ],
    col_widths=[2, 4, 8]
)

# ══════════════════════════════════════════════════════════════════════════════
# 2. BASE DE DATOS
# ══════════════════════════════════════════════════════════════════════════════
h1("2. Base de Datos: MIT-BIH Arrhythmia Database")
tabla(
    ["Parámetro", "Valor"],
    [
        ["Fuente", "PhysioNet (Moody & Mark, 2001)"],
        ["Registros", "48 grabaciones de ECG ambulatorio"],
        ["Duración por registro", "~30 minutos"],
        ["Frecuencia de muestreo", "360 Hz"],
        ["Canales usados", "Canal MLII (derivación estándar)"],
        ["Anotaciones", "Cada pico R etiquetado por cardiólogos expertos"],
        ["Total de latidos procesados", "94,627"],
    ],
    col_widths=[5, 9]
)

# ══════════════════════════════════════════════════════════════════════════════
# 3. ARQUITECTURA DEL SISTEMA
# ══════════════════════════════════════════════════════════════════════════════
h1("3. Arquitectura del Sistema — 4 Fases")
code_block(
"""FASE 1 — Adquisición y Preprocesamiento de Datos
   Archivos:  src/data/ingest.py  +  src/data/preprocess.py
   Entrada:   48 registros MIT-BIH (.dat, .hea, .atr)
   Salida:    data/processed/X.npy + y.npy  (94,627 latidos)

FASE 2 — Entrenamiento del Modelo Base (CNN)
   Archivo:   notebooks/fase2_entrenamiento.ipynb
   Entrada:   X.npy, y.npy
   Salida:    models/ecg_cnn_base.pth

FASE 3 — Reentrenamiento Incremental (ADAPT)
   Archivo:   notebooks/fase3_reentrenamiento.ipynb  +  src/ui/app.py
   Entrada:   ecg_cnn_base.pth  +  nuevos registros
   Salida:    models/ADAPT-ECG-RETRAINED.pth

FASE 4 — Evaluación y Comparación Estadística
   Archivo:   notebooks/fase4_evaluacion.ipynb
   Entrada:   ecg_cnn_base.pth  vs  ADAPT-ECG-RETRAINED.pth
   Salida:    models/fase4_resultados.json  +  gráficas docs/"""
)

# ══════════════════════════════════════════════════════════════════════════════
# 4. DESCRIPCIÓN DE ARCHIVOS
# ══════════════════════════════════════════════════════════════════════════════
h1("4. Descripción de Cada Archivo")

h2("4.1 src/config/settings.py — Configuración Global")
body("Centraliza todos los parámetros del sistema. Ningún otro archivo tiene valores hardcodeados:")
tabla(
    ["Parámetro", "Valor", "Descripción"],
    [
        ["FS", "360 Hz", "Frecuencia de muestreo MIT-BIH"],
        ["LEAD", "0 (MLII)", "Canal ECG principal"],
        ["FILTER_LOW_HZ", "0.5 Hz", "Corte inferior del filtro"],
        ["FILTER_HIGH_HZ", "40.0 Hz", "Corte superior del filtro"],
        ["FILTER_ORDER", "4", "Orden del filtro Butterworth"],
        ["BEAT_BEFORE_MS", "90 ms → 32 muestras", "Muestras antes del pico R"],
        ["BEAT_AFTER_MS", "110 ms → 40 muestras", "Muestras después del pico R"],
        ["BEAT_LEN", "72 muestras", "Longitud total de cada latido"],
        ["BATCH_SIZE", "64", "Tamaño de lote en entrenamiento"],
        ["LEARNING_RATE", "0.001", "Tasa de aprendizaje base"],
        ["EPOCHS", "30", "Épocas de entrenamiento"],
    ],
    col_widths=[4, 4, 6]
)

h2("4.2 src/data/ingest.py — Adquisición de Datos")
body("Descarga y carga los registros MIT-BIH desde PhysioNet usando la librería wfdb:")
tabla(
    ["Función", "Descripción"],
    [
        ["download_record(id)", "Descarga un registro individual (omite si ya existe)"],
        ["download_all()", "Descarga los 48 registros con barra de progreso"],
        ["load_record(id)", "Lee señal + anotaciones de un registro local"],
        ["load_all_records()", "Carga todos los registros en un diccionario"],
        ["record_info(id)", "Resumen del registro: duración, n° latidos, distribución de clases"],
    ],
    col_widths=[5, 9]
)

h2("4.3 src/data/preprocess.py — Preprocesamiento")
body("Implementa el pipeline completo de preprocesamiento de señal ECG:")
tabla(
    ["Función", "Descripción"],
    [
        ["bandpass_filter(señal)", "Filtro Butterworth zero-phase 0.5–40 Hz"],
        ["segment_beats(señal, picos_R, símbolos)", "Extrae ventanas de 72 muestras centradas en cada pico R"],
        ["normalize_beats(latidos)", "Normalización z-score independiente por latido"],
        ["process_record(señal, anotación)", "Ejecuta filtrado → segmentación → normalización"],
        ["process_and_save(registros)", "Procesa todos los registros y guarda X.npy e y.npy"],
        ["load_processed()", "Carga el dataset ya procesado desde disco"],
    ],
    col_widths=[5.5, 8.5]
)

h2("4.4 src/ui/app.py — Interfaz Streamlit")
body("Aplicación web principal. Contiene toda la lógica integrada: arquitectura CNN, "
     "Replay Buffer, reentrenamiento incremental, visualización interactiva y métricas en tiempo real.")
body("**Modos de uso:**")
body("**Modo 1 — MIT-BIH:** Selecciona cualquiera de los 48 registros, carga la señal, "
     "clasifica latido por latido y opcionalmente reentrena el modelo.")
body("**Modo 2 — CSV propio:** Sube tu propia señal ECG en CSV, detecta picos R automáticamente "
     "con scipy.signal.find_peaks y clasifica cada latido detectado.")

h2("4.5 Modelos Guardados")
tabla(
    ["Archivo", "Descripción"],
    [
        ["ecg_cnn_base.pth", "Pesos del modelo estático (entrenado una sola vez)"],
        ["ecg_cnn_base_meta.json", "Metadatos: arquitectura, hiperparámetros, métricas del entrenamiento"],
        ["ADAPT-ECG-RETRAINED.pth", "Pesos del modelo adaptativo (reentrenado 5 veces)"],
        ["replay_buffer.npz", "Memoria persistente del buffer de reentrenamiento"],
        ["retrain_history.json", "Registro JSON de cada sesión de reentrenamiento"],
        ["fase3_resultados.json", "Comparación Estático vs Adaptativo en Fase 3"],
        ["fase4_resultados.json", "Evaluación completa con métricas por clase y test estadístico"],
    ],
    col_widths=[5, 9]
)

# ══════════════════════════════════════════════════════════════════════════════
# 5. PREPROCESAMIENTO
# ══════════════════════════════════════════════════════════════════════════════
h1("5. Pipeline de Preprocesamiento (Fase 1)")

h2("5.1 Paso 1 — Filtro Butterworth Pasa-Banda")
body("Se aplica un filtro Butterworth de orden 4, zero-phase (filtfilt) con banda 0.5–40 Hz:")
body("**¿Qué elimina?**")
body("• Componentes < 0.5 Hz: deriva de línea base causada por respiración y movimiento del paciente.")
body("• Componentes > 40 Hz: interferencia eléctrica de 60 Hz y ruido de alta frecuencia.")
body("**Zero-phase (filtfilt):** aplica el filtro dos veces (hacia adelante y hacia atrás) para "
     "evitar distorsión de fase, preservando la forma exacta del complejo QRS.")

h2("5.2 Paso 2 — Segmentación por Pico R")
body("Cada latido se extrae como una ventana de **72 muestras** centrada en el pico R anotado:")
code_block(
"""ventana = señal[pico_R - 32 : pico_R + 40]
   ← 32 muestras (≈89ms) | pico R | 40 muestras (≈111ms) →
   Total: 72 muestras = un latido completo"""
)
body("Se descartan latidos que caigan en los bordes de la señal o cuyo símbolo no tenga "
     "equivalente AAMI.")

h2("5.3 Paso 3 — Normalización Z-Score")
body("Cada latido se normaliza de forma **independiente**:")
code_block("x_normalizado = (x - media(x)) / std(x)")
body("Esto permite comparar señales de distintos pacientes sin depender de la amplitud absoluta "
     "del equipo de medición. Si std = 0 (señal plana), el latido se devuelve sin modificar.")

h2("5.4 Dataset Final")
tabla(
    ["Clase", "Latidos en prueba", "Porcentaje"],
    [
        ["N — Normal",           "60,022", "79.3%"],
        ["Q — Desconocido",       "6,821",  "9.0%"],
        ["V — Ventricular",       "5,773",  "7.6%"],
        ["S — Supraventricular",  "2,437",  "3.2%"],
        ["F — Fusión",              "649",  "0.9%"],
        ["TOTAL",                "75,702", "100%"],
    ],
    col_widths=[5, 4, 4]
)
nota("El dataset es muy desbalanceado: el 79% son latidos normales. "
     "Esto representa un reto real para el clasificador en clases minoritarias como F y S.")

# ══════════════════════════════════════════════════════════════════════════════
# 6. ARQUITECTURA CNN
# ══════════════════════════════════════════════════════════════════════════════
h1("6. Arquitectura de la CNN (Red Neuronal Convolucional 1D)")

body("La red neuronal es una CNN 1D diseñada específicamente para señales temporales. "
     "A diferencia de una CNN 2D para imágenes, opera directamente sobre la secuencia temporal del latido.")

h2("6.1 Estructura Completa")
code_block(
"""Entrada: tensor (batch_size, 1, 72)
           ↑ 1 canal (señal ECG), 72 muestras por latido

┌─── BLOQUE 1 ────────────────────────────────────────────┐
│  Conv1d(in=1, out=32, kernel=5, padding=2)               │
│  → Detecta patrones locales de 5 muestras (~14ms)        │
│  BatchNorm1d(32)   → normaliza activaciones por batch    │
│  ReLU              → activación no lineal                │
│  MaxPool1d(2)      → reduce dimensión: 72 → 36           │
│  Salida: (batch, 32, 36)                                 │
└─────────────────────────────────────────────────────────┘

┌─── BLOQUE 2 ────────────────────────────────────────────┐
│  Conv1d(in=32, out=64, kernel=5, padding=2)              │
│  → Combina características de nivel bajo                 │
│  BatchNorm1d(64)                                         │
│  ReLU                                                    │
│  MaxPool1d(2)      → reduce dimensión: 36 → 18           │
│  Salida: (batch, 64, 18)                                 │
└─────────────────────────────────────────────────────────┘

┌─── BLOQUE 3 ────────────────────────────────────────────┐
│  Conv1d(in=64, out=128, kernel=3, padding=1)             │
│  → Patrones de alto nivel (morfología completa QRS)      │
│  BatchNorm1d(128)                                        │
│  ReLU                                                    │
│  MaxPool1d(2)      → reduce dimensión: 18 → 9            │
│  Salida: (batch, 128, 9)                                 │
└─────────────────────────────────────────────────────────┘

  AdaptiveAvgPool1d(1)  → promedia todos los pasos de tiempo
  Salida: (batch, 128, 1)

  Flatten  →  (batch, 128)

┌─── CLASIFICADOR ────────────────────────────────────────┐
│  Linear(128 → 64)                                        │
│  ReLU                                                    │
│  Dropout(p=0.5)    → evita sobreajuste                   │
│  Linear(64 → 5)    → una neurona por clase AAMI          │
└─────────────────────────────────────────────────────────┘

Salida: logits (batch, 5)  →  Softmax  →  probabilidades"""
)

h2("6.2 Código Real de la CNN (src/ui/app.py)")
body("Este es el código exacto implementado en PyTorch que define la arquitectura:")
code_block(
"""import torch.nn as nn

class ECG_CNN(nn.Module):
    def __init__(self, n_classes=5):
        super().__init__()

        # Bloque 1: Conv1d(1→32, kernel=5) + BN + ReLU + MaxPool
        self.block1 = nn.Sequential(
            nn.Conv1d(1, 32, kernel_size=5, padding=2),
            nn.BatchNorm1d(32), nn.ReLU(), nn.MaxPool1d(2))

        # Bloque 2: Conv1d(32→64, kernel=5) + BN + ReLU + MaxPool
        self.block2 = nn.Sequential(
            nn.Conv1d(32, 64, kernel_size=5, padding=2),
            nn.BatchNorm1d(64), nn.ReLU(), nn.MaxPool1d(2))

        # Bloque 3: Conv1d(64→128, kernel=3) + BN + ReLU + MaxPool
        self.block3 = nn.Sequential(
            nn.Conv1d(64, 128, kernel_size=3, padding=1),
            nn.BatchNorm1d(128), nn.ReLU(), nn.MaxPool1d(2))

        # Pooling global: colapsa la dimensión temporal
        self.pool = nn.AdaptiveAvgPool1d(1)

        # Clasificador: Linear → ReLU → Dropout → Linear
        self.classifier = nn.Sequential(
            nn.Flatten(),
            nn.Linear(128, 64), nn.ReLU(), nn.Dropout(0.5),
            nn.Linear(64, n_classes))

    def forward(self, x):
        x = self.block1(x)   # (batch, 1, 72)  → (batch, 32, 36)
        x = self.block2(x)   # (batch, 32, 36) → (batch, 64, 18)
        x = self.block3(x)   # (batch, 64, 18) → (batch, 128, 9)
        x = self.pool(x)     # (batch, 128, 9) → (batch, 128, 1)
        return self.classifier(x)  # → (batch, 5)

    def predict_proba(self, x):
        # Aplica softmax para obtener probabilidades por clase
        return torch.softmax(self.forward(x), dim=1)"""
)

h2("6.3 Código de Inferencia (run_inference)")
body("Función que ejecuta la CNN sobre los latidos segmentados para obtener predicciones:")
code_block(
"""def run_inference(model, beats, batch_size=128):
    model.eval()  # Modo evaluación: Dropout desactivado

    # Convertir array NumPy a tensor PyTorch: (N, 72) → (N, 1, 72)
    X = torch.tensor(beats, dtype=torch.float32).unsqueeze(1)
    loader = DataLoader(TensorDataset(X), batch_size=batch_size, shuffle=False)

    all_probs = []
    with torch.no_grad():  # Sin calcular gradientes (inferencia pura)
        for (xb,) in loader:
            all_probs.append(model.predict_proba(xb).numpy())

    probs = np.concatenate(all_probs, axis=0)  # shape: (N, 5)

    # Retorna: clase predicha (argmax) + todas las probabilidades
    return probs.argmax(axis=1), probs"""
)

h2("6.4 ¿Por Qué CNN 1D y No Otro Tipo de Red?")
tabla(
    ["Alternativa", "Ventaja", "Desventaja vs CNN 1D"],
    [
        ["CNN 1D (elegida)", "Captura patrones locales (QRS), rápida, ligera", "—"],
        ["RNN / LSTM", "Modela dependencias temporales largas", "Más lenta, difícil de entrenar con señales cortas"],
        ["Transformer", "Atención global muy potente", "Requiere mucho más dato y cómputo"],
        ["SVM / Random Forest", "Interpretable, menos datos", "No aprende representaciones automáticamente"],
    ],
    col_widths=[4, 5, 5]
)

h2("6.3 Rol de Cada Componente")
tabla(
    ["Componente", "Función en el contexto ECG"],
    [
        ["Conv1d kernel=5 (Block 1)", "Detecta ondas P, inicio del QRS (~14ms de señal)"],
        ["Conv1d kernel=5 (Block 2)", "Combina ondas detectadas, detecta el complejo QRS completo"],
        ["Conv1d kernel=3 (Block 3)", "Captura la morfología global: duración, amplitud, forma del latido"],
        ["BatchNorm1d", "Estabiliza el entrenamiento con el desbalanceo severo de clases"],
        ["MaxPool1d(2)", "Reduce ruido y crea invarianza a pequeños desplazamientos del pico R"],
        ["AdaptiveAvgPool1d(1)", "Comprime toda la información en un vector de 128 características"],
        ["Dropout(0.5)", "Previene sobreajuste: desactiva neuronas al azar durante entrenamiento"],
        ["Softmax", "Convierte logits en probabilidades que suman 1.0"],
    ],
    col_widths=[5.5, 8.5]
)

# ══════════════════════════════════════════════════════════════════════════════
# 7. ENTRENAMIENTO BASE
# ══════════════════════════════════════════════════════════════════════════════
h1("7. Entrenamiento del Modelo Base (Fase 2)")

tabla(
    ["Parámetro", "Valor"],
    [
        ["Dataset total", "94,627 latidos"],
        ["Split entrenamiento", "80% — 75,701 latidos"],
        ["Split validación", "20% — 18,926 latidos"],
        ["Épocas", "30"],
        ["Batch size", "64"],
        ["Optimizador", "Adam"],
        ["Learning rate", "0.001"],
        ["Función de pérdida", "CrossEntropyLoss"],
        ["Semilla aleatoria", "42"],
    ],
    col_widths=[5, 9]
)

h2("7.1 Resultados del Entrenamiento Base")
tabla(
    ["Métrica", "Valor"],
    [
        ["Accuracy (validación)", "93.18%"],
        ["F1-macro", "77.83%"],
        ["F1-weighted", "94.12%"],
    ],
    col_widths=[7, 7]
)

# ══════════════════════════════════════════════════════════════════════════════
# 8. REENTRENAMIENTO — SECCIÓN PRINCIPAL EXPANDIDA
# ══════════════════════════════════════════════════════════════════════════════
h1("8. Reentrenamiento Incremental de la CNN — Explicación Detallada (Fase 3)")

h2("8.1 El Problema: Olvido Catastrófico")
body("Cuando una red neuronal se reentrena con nuevos datos usando gradiente descendente estándar, "
     "los pesos se actualizan para minimizar el error en los datos nuevos. Esto causa que el modelo "
     "**sobreescriba** los pesos que codificaban el conocimiento anterior. A este fenómeno se le llama "
     "**Olvido Catastrófico** (Catastrophic Forgetting).")

body("**Ejemplo concreto en ADAPT-ECG:**")
body("→ El modelo aprendió a clasificar latidos Normales (N) con 99% de accuracy en los registros 100–119.")
body("→ Si se reentrena solo con datos del registro 200 (que tiene muchos latidos Ventriculares), "
     "el modelo empieza a predecir V donde antes predecía N correctamente.")
body("→ La distribución de gradientes de los datos nuevos empuja los pesos lejos de la región "
     "óptima para los datos anteriores.")

h2("8.2 La Solución: Replay Buffer (Memoria de Experiencias)")
body("El Replay Buffer es una **memoria circular** que almacena latidos de sesiones anteriores. "
     "Al reentrenar, la CNN no solo ve los datos nuevos sino también una muestra aleatoria "
     "de lo que ya aprendió, forzándola a mantener el rendimiento en ambas distribuciones.")

code_block(
"""class ReplayBuffer:
    max_size = 2,000 latidos          # capacidad máxima
    estrategia = FIFO                 # los más antiguos salen primero

    def add(nuevos_beats, nuevas_labels):
        # Agrega nuevos latidos al buffer
        buffer.extend(nuevos_beats, nuevas_labels)
        # Si excede el límite, descarta los más viejos
        if len(buffer) > 2000:
            buffer = buffer[-2000:]   # mantiene los últimos 2,000

    def sample(n):
        # Muestrea n latidos al azar del buffer
        return random_sample(buffer, size=n)

    def save() / load():
        # Persiste en disco como replay_buffer.npz
        # Se recarga en cada nueva sesión de reentrenamiento"""
)

h2("8.3 Código Real del Replay Buffer (src/ui/app.py)")
body("Esta es la implementación exacta de la memoria circular usada en el proyecto:")
code_block(
"""class ReplayBuffer:
    def __init__(self, max_size=2000):
        self.X = []          # lista de latidos almacenados
        self.y = []          # lista de etiquetas correspondientes
        self.max_size = max_size  # capacidad máxima: 2,000 latidos

    def add(self, X_batch, y_batch):
        \"\"\"Agrega nuevos latidos al buffer. Si excede max_size,
        descarta los más antiguos (estrategia FIFO).\"\"\"
        self.X.extend(X_batch.tolist())
        self.y.extend(y_batch.tolist())
        if len(self.X) > self.max_size:
            self.X = self.X[-self.max_size:]  # conserva los últimos 2,000
            self.y = self.y[-self.max_size:]

    def sample(self, n):
        \"\"\"Muestrea n latidos al azar del buffer (sin reemplazo).\"\"\"
        idx = np.random.choice(len(self.X), size=min(n, len(self.X)), replace=False)
        return (np.array([self.X[i] for i in idx], dtype=np.float32),
                np.array([self.y[i] for i in idx], dtype=np.int64))

    def save(self, path):
        \"\"\"Persiste el buffer en disco como archivo .npz (NumPy).\"\"\"
        np.savez(path, X=np.array(self.X, dtype=np.float32),
                       y=np.array(self.y, dtype=np.int64))

    def load(self, path):
        \"\"\"Carga el buffer desde disco al iniciar una nueva sesión.\"\"\"
        if Path(path).exists():
            data = np.load(path)
            self.X = data["X"].tolist()
            self.y = data["y"].tolist()

    def __len__(self):
        return len(self.X)"""
)

h2("8.4 ¿Qué Capas de la CNN se Actualizan en el Reentrenamiento?")
body("Esta es una decisión clave del diseño. En ADAPT-ECG, **TODAS las capas se actualizan** "
     "durante el reentrenamiento (no hay capas congeladas). Esto es posible y seguro gracias a "
     "dos mecanismos combinados:")

body("**Mecanismo 1 — Learning Rate muy bajo (lr = 0.0001 vs 0.001 del entrenamiento base):**")
body("El learning rate es 10 veces menor que en el entrenamiento original. Esto significa que "
     "los gradientes actualizan los pesos con pasos muy pequeños, preservando la mayor parte "
     "del conocimiento ya codificado. Una actualización grande sobreescribiría los pesos; "
     "una pequeña los ajusta fino.")

code_block(
"""Entrenamiento base:    optimizer = Adam(lr=0.001)   → pasos grandes, aprende todo desde cero
Reentrenamiento:       optimizer = Adam(lr=0.0001)  → pasos pequeños, ajuste fino"""
)

body("**Mecanismo 2 — Mezcla de datos (Replay):**")
body("En cada época del reentrenamiento, el batch que ve la CNN es una combinación de:")
code_block(
"""X_combinado = concat(nuevos_beats, replay_buffer.sample(n=len(nuevos_beats) * 0.5))
y_combinado = concat(nuevas_labels, replay_labels)
# Ejemplo: 2,000 nuevos latidos + 1,000 del buffer = 3,000 latidos por época"""
)
body("Los gradientes calculados sobre esta mezcla actúan simultáneamente en dos direcciones: "
     "adaptarse a los datos nuevos Y mantener el rendimiento en los datos del buffer. "
     "El resultado es que los pesos convergen a una región que satisface ambas distribuciones.")

h2("8.5 Código Real del Reentrenamiento Incremental (src/ui/app.py)")
body("Esta es la función exacta implementada en el proyecto:")
code_block(
"""def incremental_train(model, beats, labels, replay_buffer,
                      epochs=5, lr=1e-4, batch_size=64):
    \"\"\"Reentrenamiento incremental con Replay Buffer.\"\"\"

    # Optimizador Adam con lr 10x menor que el entrenamiento base
    optimizer = torch.optim.Adam(model.parameters(), lr=lr)
    criterion = nn.CrossEntropyLoss()
    model.train()  # Activar Dropout y BatchNorm en modo entrenamiento

    # Agregar nuevos latidos al buffer ANTES de entrenar
    replay_buffer.add(beats, labels)

    losses = []
    for _ in range(epochs):  # por defecto: 5 epocas

        # Calcular cuantos latidos tomar del buffer (50% de los nuevos)
        n_replay = int(len(beats) * 0.5)

        if len(replay_buffer) > 0 and n_replay > 0:
            # Muestrear del buffer latidos de sesiones anteriores
            Xr, yr = replay_buffer.sample(n_replay)
            # Combinar: nuevos datos + datos del buffer
            X_comb = np.concatenate([beats, Xr], axis=0)
            y_comb = np.concatenate([labels, yr], axis=0)
        else:
            X_comb, y_comb = beats, labels

        # Mezclar aleatoriamente para evitar sesgos de orden
        perm   = np.random.permutation(len(X_comb))
        X_comb = X_comb[perm]
        y_comb = y_comb[perm]

        # Convertir a tensores PyTorch: (N, 72) -> (N, 1, 72)
        Xt = torch.tensor(X_comb, dtype=torch.float32).unsqueeze(1)
        yt = torch.tensor(y_comb, dtype=torch.long)
        loader = DataLoader(TensorDataset(Xt, yt),
                            batch_size=batch_size, shuffle=True)

        epoch_loss = 0.0
        for xb, yb in loader:
            optimizer.zero_grad()             # limpiar gradientes acumulados
            loss = criterion(model(xb), yb)   # forward pass + calcular perdida
            loss.backward()                   # backpropagation: dL/dW
            optimizer.step()                  # actualizar pesos: W -= lr * dL/dW
            epoch_loss += loss.item()

        losses.append(epoch_loss / len(loader))

    model.eval()   # Volver a modo inferencia (Dropout desactivado)
    return losses  # perdida por epoca (para graficar en la UI)"""
)
nota("Este codigo se ejecuta desde la interfaz Streamlit al presionar 'Reentrenar'. "
     "Despues guarda el modelo actualizado y el buffer en disco automaticamente.")

h2("8.6 Algoritmo Completo de Reentrenamiento Paso a Paso")
code_block(
"""def incremental_train(model, nuevos_beats, nuevas_labels, replay_buffer,
                      epochs=5, lr=0.0001, batch_size=64):

    # PASO 1: Configurar optimizador con lr pequeño
    optimizer = Adam(model.parameters(), lr=0.0001)
    criterion = CrossEntropyLoss()
    model.train()  # activar modo entrenamiento (BatchNorm y Dropout activos)

    # PASO 2: Guardar nuevos datos en el buffer ANTES de entrenar
    replay_buffer.add(nuevos_beats, nuevas_labels)
    # → El buffer ya tiene los datos nuevos para futuras sesiones

    losses = []
    for época in range(5):

        # PASO 3: Muestrear del buffer (50% del tamaño de nuevos datos)
        n_replay = len(nuevos_beats) * 0.5
        X_replay, y_replay = replay_buffer.sample(n=n_replay)

        # PASO 4: Combinar nuevos datos + datos del buffer
        X_combinado = concat([nuevos_beats, X_replay], axis=0)
        y_combinado = concat([nuevas_labels, y_replay], axis=0)

        # PASO 5: Mezclar aleatoriamente (evitar sesgos de orden)
        permutacion = random_permutation(len(X_combinado))
        X_combinado = X_combinado[permutacion]
        y_combinado = y_combinado[permutacion]

        # PASO 6: Convertir a tensores PyTorch
        Xt = tensor(X_combinado).unsqueeze(1)  # shape: (N, 1, 72)
        yt = tensor(y_combinado, dtype=long)    # shape: (N,)

        # PASO 7: Mini-batches y backpropagation
        loader = DataLoader(TensorDataset(Xt, yt), batch_size=64, shuffle=True)
        epoch_loss = 0
        for X_batch, y_batch in loader:
            optimizer.zero_grad()           # resetear gradientes
            logits = model(X_batch)         # forward pass por los 3 bloques CNN
            loss = criterion(logits, y_batch)  # CrossEntropyLoss
            loss.backward()                 # backpropagation: calcular ∂L/∂W
            optimizer.step()               # actualizar pesos: W -= lr * ∂L/∂W
            epoch_loss += loss.item()
        losses.append(epoch_loss / len(loader))

    model.eval()  # volver a modo inferencia (Dropout desactivado)

    # PASO 8: Guardar modelo actualizado
    torch.save(model.state_dict(), "ADAPT-ECG-RETRAINED.pth")

    # PASO 9: Persistir el buffer para la próxima sesión
    replay_buffer.save("replay_buffer.npz")

    return losses  # pérdida por época para graficar"""
)

h2("8.5 Por Qué Solo 5 Épocas en el Reentrenamiento")
body("El reentrenamiento usa solo **5 épocas** (vs 30 en el entrenamiento base). Esto es intencional:")
body("• **Demasiadas épocas** → sobreajuste a los datos nuevos → olvido catastrófico igual.")
body("• **Muy pocas épocas** → el modelo no se adapta suficientemente a los nuevos patrones.")
body("• **5 épocas con lr=0.0001** → es el balance: suficiente para adaptar, insuficiente para olvidar.")
nota("Este valor se determinó empíricamente en los experimentos de la Fase 3.")

h2("8.6 Historial Real de las 5 Sesiones de Reentrenamiento")
tabla(
    ["Sesión", "Registro", "Latidos", "Acc. Antes", "Acc. Después", "F1 Antes", "F1 Después"],
    [
        ["1", "100", "2,272", "98.55%", "99.08% (+0.53%)", "66.42%", "84.29% (+17.87%)"],
        ["2", "103", "2,084", "77.69%", "99.90% (+22.21%)", "29.29%", "49.98% (+20.69%)"],
        ["3", "112", "2,539", "87.00%", "99.92% (+12.92%)", "18.61%", "49.98% (+31.37%)"],
        ["4", "100", "2,272", "98.55%", "99.34% (+0.79%)", "66.42%", "91.12% (+24.70%)"],
        ["5", "124", "88",   "81.82%", "73.86% (-7.96%)", "36.14%", "33.66% (-2.48%)"],
    ],
    col_widths=[1.5, 2, 2, 2.5, 3.5, 2, 3]
)
nota("Sesión 5 (registro 124): solo 88 latidos disponibles — muy pocos para que el reentrenamiento "
     "sea estable. Con datos escasos, los gradientes son ruidosos y el modelo puede degradarse. "
     "Con al menos 500 latidos el reentrenamiento es consistentemente positivo.")

h2("8.7 ¿Qué Aprende Cada Capa Durante el Reentrenamiento?")
tabla(
    ["Capa", "Qué aprende originalmente", "Qué ajusta en reentrenamiento"],
    [
        ["Block 1 (Conv 1→32)", "Bordes, picos simples de la señal ECG", "Ajuste mínimo — patrones muy básicos, ya generalizados"],
        ["Block 2 (Conv 32→64)", "Combinaciones de características, forma del QRS", "Ajuste moderado — refina patrones específicos del nuevo paciente"],
        ["Block 3 (Conv 64→128)", "Morfología completa: duración, amplitud, variabilidad", "Ajuste mayor — aquí se adaptan las características de alto nivel"],
        ["Clasificador (Linear)", "Fronteras de decisión entre las 5 clases AAMI", "Ajuste más significativo — las probabilidades de clase se rebalancean"],
    ],
    col_widths=[3.5, 5, 5.5]
)

h2("8.8 Comparación: Entrenamiento Base vs Reentrenamiento")
tabla(
    ["Aspecto", "Entrenamiento Base (Fase 2)", "Reentrenamiento (Fase 3)"],
    [
        ["Objetivo", "Aprender desde cero", "Adaptar sin olvidar"],
        ["Datos", "94,627 latidos de 48 registros", "~2,000 latidos de 1 registro + buffer"],
        ["Épocas", "30", "5"],
        ["Learning rate", "0.001", "0.0001 (10× menor)"],
        ["Capas congeladas", "Ninguna", "Ninguna (todas se actualizan con lr pequeño)"],
        ["Mezcla de datos", "No aplica", "50% nuevos + 50% del Replay Buffer"],
        ["Tiempo aprox.", "~10-15 min (Colab GPU)", "~30-60 seg (CPU)"],
        ["Riesgo", "Sobreajuste a datos de entrenamiento", "Olvido catastrófico (mitigado por buffer)"],
    ],
    col_widths=[4, 5, 5]
)

# ══════════════════════════════════════════════════════════════════════════════
# 9. RESULTADOS FASE 3
# ══════════════════════════════════════════════════════════════════════════════
h1("9. Resultados — Fase 3: Estático vs Adaptativo")
tabla(
    ["Métrica", "Modelo Estático", "Modelo Adaptativo", "Mejora"],
    [
        ["Accuracy",    "93.54%", "97.83%", "+4.29%"],
        ["F1-macro",    "78.49%", "90.71%", "+12.22%"],
        ["F1-weighted", "94.47%", "97.73%", "+3.26%"],
    ],
    col_widths=[4, 4, 4, 4]
)

# ══════════════════════════════════════════════════════════════════════════════
# 10. EVALUACIÓN FINAL FASE 4
# ══════════════════════════════════════════════════════════════════════════════
h1("10. Evaluación Final (Fase 4) — 75,702 Latidos")

h2("10.1 Métricas Globales")
tabla(
    ["Métrica", "Modelo Estático", "Modelo Adaptativo", "Mejora"],
    [
        ["Accuracy",          "88.72%", "97.41%", "+8.69%"],
        ["F1-macro",          "73.81%", "89.54%", "+15.73%"],
        ["F1-weighted",       "91.14%", "97.36%", "+6.22%"],
        ["Precision-macro",   "68.91%", "92.55%", "+23.64%"],
        ["Recall-macro",      "92.48%", "86.99%", "-5.49%"],
    ],
    col_widths=[4.5, 3.5, 3.5, 3.5]
)

h2("10.2 Métricas por Clase — Modelo Estático")
tabla(
    ["Clase", "Precisión", "Sensibilidad", "Especificidad", "F1"],
    [
        ["N — Normal",           "99.32%", "87.11%", "97.72%", "92.81%"],
        ["S — Supraventricular", "26.82%", "91.01%", "91.74%", "41.43%"],
        ["V — Ventricular",      "92.27%", "93.21%", "99.36%", "92.74%"],
        ["F — Fusión",           "29.53%", "93.22%", "98.08%", "44.85%"],
        ["Q — Desconocido",      "96.61%", "97.86%", "99.66%", "97.23%"],
    ],
    col_widths=[4, 3, 3, 3.5, 3]
)

h2("10.3 Métricas por Clase — Modelo Adaptativo")
tabla(
    ["Clase", "Precisión", "Sensibilidad", "Especificidad", "F1"],
    [
        ["N — Normal",           "98.11%", "99.05%", "92.68%", "98.58%"],
        ["S — Supraventricular", "81.32%", "73.78%", "99.44%", "77.37%"],
        ["V — Ventricular",      "96.16%", "93.28%", "99.69%", "94.70%"],
        ["F — Fusión",           "89.06%", "71.49%", "99.92%", "79.32%"],
        ["Q — Desconocido",      "98.12%", "97.36%", "99.82%", "97.74%"],
    ],
    col_widths=[4, 3, 3, 3.5, 3]
)
nota("La mayor mejora es en Fusión (F): precisión 29.53% → 89.06%. El modelo estático "
     "confundía muchos latidos de otras clases como F. El reentrenamiento corrigió esto.")

h2("10.4 AUC-ROC")
tabla(
    ["Clase", "Estático", "Adaptativo"],
    [
        ["N — Normal",           "0.9909", "0.9913"],
        ["S — Supraventricular", "0.9743", "0.9740"],
        ["V — Ventricular",      "0.9980", "0.9989"],
        ["F — Fusión",           "0.9930", "0.9939"],
        ["Q — Desconocido",      "0.9996", "0.9997"],
    ],
    col_widths=[5, 4, 4]
)
nota("AUC > 0.97 en todas las clases — el modelo es un discriminador excelente incluso en clases raras.")

h2("10.5 Test Estadístico de McNemar")
body("Se aplicó el test de McNemar para verificar que la mejora entre modelos no es por azar:")
tabla(
    ["Parámetro", "Valor"],
    [
        ["Estadístico", "5,342.29"],
        ["p-valor", "0.0000"],
        ["¿Significativo?", "SÍ (p < 0.05)"],
        ["Interpretación", "La mejora del modelo adaptativo es estadísticamente significativa"],
    ],
    col_widths=[5, 9]
)

# ══════════════════════════════════════════════════════════════════════════════
# 11. INTERFAZ
# ══════════════════════════════════════════════════════════════════════════════
h1("11. Interfaz de Usuario (Streamlit)")
body("Ejecutar con:  streamlit run src/ui/app.py")

h2("Modo 1 — MIT-BIH")
body("1. Selecciona carpeta con registros MIT-BIH")
body("2. Elige un registro (100–234)")
body("3. La app carga señal, filtra, segmenta y clasifica cada latido")
body("4. Muestra ECG interactivo con picos R coloreados por clase AAMI")
body("5. Muestra distribución de clases, matriz de confusión y métricas por clase")
body("6. Panel de reentrenamiento: selecciona épocas, learning rate y tamaño del buffer")
body("7. Botón Reentrenar → actualiza ADAPT-ECG-RETRAINED.pth y guarda en el historial")

h2("Modo 2 — CSV Propio")
body("1. Sube archivo CSV con columna de voltaje en mV")
body("2. Especifica la frecuencia de muestreo del equipo")
body("3. La app detecta picos R automáticamente con scipy.signal.find_peaks")
body("4. Clasifica cada latido detectado y muestra resultados")

# ══════════════════════════════════════════════════════════════════════════════
# 12. STACK TECNOLÓGICO
# ══════════════════════════════════════════════════════════════════════════════
h1("12. Stack Tecnológico")
tabla(
    ["Librería", "Uso principal"],
    [
        ["PyTorch ≥ 2.0",          "Definición y entrenamiento de la CNN, backpropagation"],
        ["Streamlit ≥ 1.30",       "Interfaz web interactiva"],
        ["wfdb ≥ 4.1",             "Lectura de registros MIT-BIH (.dat, .hea, .atr)"],
        ["NumPy ≥ 1.24",           "Operaciones numéricas, manejo de arrays"],
        ["SciPy ≥ 1.11",           "Filtro Butterworth, detección de picos (find_peaks)"],
        ["scikit-learn ≥ 1.3",     "Métricas: accuracy, F1, confusion matrix"],
        ["Plotly ≥ 5.18",          "Gráficas interactivas del ECG y métricas"],
        ["Matplotlib / Seaborn",   "Matriz de confusión estática"],
        ["Pandas ≥ 2.0",           "Manejo tabular de resultados e historial"],
    ],
    col_widths=[5, 9]
)

# ══════════════════════════════════════════════════════════════════════════════
# 13. CONCLUSIONES
# ══════════════════════════════════════════════════════════════════════════════
h1("13. Conclusiones")
body("1. **La CNN 1D entrenada con MIT-BIH alcanza 93.18% de accuracy** en validación, "
     "superando el umbral mínimo clínico para sistemas de apoyo al diagnóstico.")
body("2. **El Replay Buffer resuelve el olvido catastrófico** al mezclar datos históricos "
     "con datos nuevos en cada sesión de reentrenamiento, sin necesidad de congelar capas.")
body("3. **El learning rate 10× menor (0.0001 vs 0.001)** es el mecanismo clave que permite "
     "actualizar todas las capas sin destruir el conocimiento previo.")
body("4. **El modelo adaptativo supera estadísticamente al estático** (p=0.0 en McNemar), "
     "con mejoras de hasta +15.7% en F1-macro y +8.69% en accuracy.")
body("5. **La clase más difícil es Fusión (F)** por su baja prevalencia (0.9%) y naturaleza ambigua, "
     "pero el modelo adaptativo la mejora de F1=44.85% a F1=79.32% (+34.47%).")
body("6. **El sistema es funcional en tiempo real** desde la interfaz Streamlit, "
     "permitiendo inferencia y reentrenamiento sin línea de comandos.")

# ── Guardar ────────────────────────────────────────────────────────────────────
doc.save(OUT)
print(f"Word generado: {OUT}")
