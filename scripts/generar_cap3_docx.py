"""
Genera docs/Cap3_Desarrollo_Resultados.docx
SOLO sección de Desarrollo (los Resultados se añaden al terminar el reentrenamiento).
Modelo activo: ECG_ResNet_SE, BEAT_LEN=128.
Formato ITT Residencia Profesional:
  - Times New Roman 12 pt
  - Interlineado 1.5 (21 pt)
  - Márgenes 2.5 cm
  - Títulos: mayúscula, izquierda, negrita, 14 pt
  - Texto justificado, sin sangría
  - Un renglón entre párrafos
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
OUT  = DOCS / "Cap3_Desarrollo_Resultados.docx"

LINE_SPACING = Pt(21)   # 1.5 × 14 pt ≈ 21 pt para texto 12 pt

# ─────────────────────────────────────────────────────────────────────────────
# Helpers de formato
# ─────────────────────────────────────────────────────────────────────────────

def _set_font(run, name="Times New Roman"):
    rPr = run._r.get_or_add_rPr()
    rf  = OxmlElement("w:rFonts")
    rf.set(qn("w:ascii"), name)
    rf.set(qn("w:hAnsi"), name)
    rPr.insert(0, rf)

def fmt(run, bold=False, italic=False, size=12, color=None):
    run.font.name   = "Times New Roman"
    run.font.size   = Pt(size)
    run.font.bold   = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color
    _set_font(run)

def _base_para(doc, align=WD_ALIGN_PARAGRAPH.JUSTIFY, sb=0, sa=6):
    p = doc.add_paragraph()
    p.alignment = align
    pf = p.paragraph_format
    pf.space_before    = Pt(sb)
    pf.space_after     = Pt(sa)
    pf.line_spacing    = LINE_SPACING
    pf.first_line_indent = Pt(0)
    return p

def txt(doc, text, bold=False, italic=False, size=12,
        align=WD_ALIGN_PARAGRAPH.JUSTIFY, sb=0, sa=6):
    p = _base_para(doc, align=align, sb=sb, sa=sa)
    if text:
        r = p.add_run(text)
        fmt(r, bold=bold, italic=italic, size=size)
    return p

def mixed(doc, parts, align=WD_ALIGN_PARAGRAPH.JUSTIFY, sb=0, sa=6):
    """parts: list of (text, bold, italic) or (text, bold, italic, size)"""
    p = _base_para(doc, align=align, sb=sb, sa=sa)
    for item in parts:
        t, b, i = item[0], item[1], item[2]
        sz = item[3] if len(item) > 3 else 12
        r = p.add_run(t)
        fmt(r, bold=b, italic=i, size=sz)
    return p

def h1(doc, text):
    """Capítulo — centrado, mayúsculas, negrita 14"""
    p = _base_para(doc, align=WD_ALIGN_PARAGRAPH.CENTER, sb=18, sa=12)
    r = p.add_run(text.upper())
    fmt(r, bold=True, size=14)

def h2(doc, text):
    """Sección — izquierda, negrita 14"""
    p = _base_para(doc, align=WD_ALIGN_PARAGRAPH.LEFT, sb=14, sa=6)
    r = p.add_run(text.upper())
    fmt(r, bold=True, size=14)

def h3(doc, text):
    """Subsección — izquierda, negrita 12"""
    p = _base_para(doc, align=WD_ALIGN_PARAGRAPH.LEFT, sb=10, sa=4)
    r = p.add_run(text)
    fmt(r, bold=True, size=12)

def h4(doc, text):
    """Sub-subsección — izquierda, negrita-cursiva 12"""
    p = _base_para(doc, align=WD_ALIGN_PARAGRAPH.LEFT, sb=8, sa=2)
    r = p.add_run(text)
    fmt(r, bold=True, italic=True, size=12)

def cap_tab(doc, text):
    p = _base_para(doc, align=WD_ALIGN_PARAGRAPH.CENTER, sb=8, sa=3)
    r = p.add_run(text)
    fmt(r, bold=True, size=11)

def cap_fig(doc, text):
    p = _base_para(doc, align=WD_ALIGN_PARAGRAPH.CENTER, sb=2, sa=10)
    r = p.add_run(text)
    fmt(r, italic=True, size=11)

def blank(doc):
    txt(doc, "", sa=0, sb=0)

def num_list(doc, items):
    """items: [(label_bold, body_text)]"""
    for label, body in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after   = Pt(4)
        p.paragraph_format.space_before  = Pt(0)
        p.paragraph_format.line_spacing  = LINE_SPACING
        p.paragraph_format.first_line_indent = Pt(0)
        r1 = p.add_run(label); fmt(r1, bold=True)
        r2 = p.add_run(body);  fmt(r2)

def bullet(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after  = Pt(3)
        p.paragraph_format.line_spacing = LINE_SPACING
        p.paragraph_format.first_line_indent = Pt(0)
        r = p.add_run(item); fmt(r)

# ── Tabla helpers ─────────────────────────────────────────────────────────────

def _shade(row, hex_c="D6E4F0"):
    for cell in row.cells:
        tc   = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd  = OxmlElement("w:shd")
        shd.set(qn("w:val"),   "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"),  hex_c)
        tcPr.append(shd)

def _cell(cell, text, bold=False, italic=False, center=True, sz=10, color=None):
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    r = p.add_run(text)
    fmt(r, bold=bold, italic=italic, size=sz, color=color)

def hdr(row, texts, bg="1F4E79"):
    _shade(row, bg)
    for i, t in enumerate(texts):
        c = row.cells[i]
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = c.paragraphs[0].add_run(t)
        fmt(r, bold=True, size=10)
        r.font.color.rgb = RGBColor(255, 255, 255)

def new_table(doc, rows, cols):
    t = doc.add_table(rows=rows, cols=cols)
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    return t

# ─────────────────────────────────────────────────────────────────────────────
# DOCUMENTO
# ─────────────────────────────────────────────────────────────────────────────
doc = Document()

for sec in doc.sections:
    sec.top_margin    = Cm(2.5)
    sec.bottom_margin = Cm(2.5)
    sec.left_margin   = Cm(2.5)
    sec.right_margin  = Cm(2.5)

# ═════════════════════════════════════════════════════════════════════════════
# CAPÍTULO 3
# ═════════════════════════════════════════════════════════════════════════════
h1(doc, "Capítulo 3\nDesarrollo y resultados")

txt(doc,
    "El presente capítulo describe el proceso completo de desarrollo del sistema "
    "inteligente con reentrenamiento continuo para la detección adaptativa de "
    "patologías cardiovasculares. El desarrollo se estructuró en cuatro fases "
    "secuenciales: adquisición y preprocesamiento de datos (Fase 1), diseño y "
    "entrenamiento del modelo neuronal base (Fase 2), implementación del módulo "
    "de reentrenamiento continuo (Fase 3) y evaluación comparativa formal (Fase 4). "
    "Para cada fase se detallan los materiales, herramientas, decisiones de diseño "
    "y métodos empleados.")

# ═════════════════════════════════════════════════════════════════════════════
# 3.1 DESARROLLO
# ═════════════════════════════════════════════════════════════════════════════
h2(doc, "3.1  Desarrollo")

# ─────────────────────────────────────────────────────────────────────────────
h3(doc, "3.1.1  Entorno de desarrollo y herramientas")
# ─────────────────────────────────────────────────────────────────────────────

txt(doc,
    "El proyecto se desarrolló en dos entornos computacionales complementarios. "
    "Las tareas de cómputo intensivo —preprocesamiento masivo, entrenamiento del "
    "modelo base y reentrenamiento continuo— se ejecutaron en Google Colaboratory "
    "con aceleración por GPU NVIDIA Tesla T4 (16 GB VRAM). Las tareas de análisis, "
    "evaluación e interfaz de usuario se realizaron en un equipo local con "
    "Windows 11, procesador Intel Core i7 y 8 GB RAM, usando Python 3.12 en "
    "entorno virtual (venv). El código fuente fue gestionado con Git y GitHub.")

blank(doc)
cap_tab(doc, "Tabla 3.1. Bibliotecas y frameworks principales del proyecto.")
t = new_table(doc, 8, 3)
hdr(t.rows[0], ["Biblioteca", "Versión", "Uso en el proyecto"])
libs = [
    ("PyTorch",          "2.x",       "Definición, entrenamiento e inferencia del modelo neuronal"),
    ("scikit-learn",     "1.x",       "Partición estratificada, métricas de evaluación"),
    ("NumPy / SciPy",    "1.x",       "Operaciones matriciales, filtrado de señales, prueba McNemar"),
    ("wfdb",             "4.x",       "Descarga y lectura de registros MIT-BIH de PhysioNet"),
    ("Streamlit",        "1.x",       "Interfaz de usuario web interactiva"),
    ("Plotly / Matplotlib", "5.x / 3.x", "Visualizaciones interactivas y estáticas"),
    ("Git + GitHub",     "—",         "Control de versiones y respaldo remoto del proyecto"),
]
for i, (lib, ver, uso) in enumerate(libs):
    row = t.rows[i + 1]
    _shade(row, "D6E4F0" if i % 2 == 0 else "FFFFFF")
    _cell(row.cells[0], lib, bold=True, center=False)
    _cell(row.cells[1], ver, center=True)
    _cell(row.cells[2], uso, center=False)

blank(doc)

# ─────────────────────────────────────────────────────────────────────────────
h3(doc, "3.1.2  Conjunto de datos: MIT-BIH Arrhythmia Database")
# ─────────────────────────────────────────────────────────────────────────────

txt(doc,
    "Se utilizó la MIT-BIH Arrhythmia Database [1], disponible públicamente en "
    "PhysioNet [2], descargada de forma programática mediante la biblioteca wfdb "
    "de Python. La base de datos comprende 48 registros de ECG ambulatorio de dos "
    "canales obtenidos de 47 pacientes distintos, con una frecuencia de muestreo "
    "de 360 Hz, resolución de 11 bits en un rango de ±5 mV y duración aproximada "
    "de 30 minutos por registro. Las anotaciones de latidos fueron revisadas de "
    "forma independiente por dos cardiólogos expertos.")

txt(doc,
    "Las etiquetas originales (más de 15 tipos distintos) fueron agrupadas en "
    "cinco clases según el estándar internacional AAMI EC57 [3], tal como se "
    "indica en la Tabla 3.2.")

blank(doc)
cap_tab(doc, "Tabla 3.2. Mapeo de etiquetas MIT-BIH a las cinco clases del estándar AAMI EC57.")
t2 = new_table(doc, 6, 4)
hdr(t2.rows[0], ["Clase", "Descripción", "Etiquetas MIT-BIH originales", "Relevancia clínica"])
aami = [
    ("N", "Normal y de unión",               "N, .",              "Ritmo sinusal normal, unión AV"),
    ("S", "Supraventricular ectópico",        "A, a, J, S, e, j", "Flutter, taquicardia supraventricular"),
    ("V", "Ventricular ectópico",             "V, E",             "Extrasístoles, taquicardia ventricular"),
    ("F", "Fusión ventricular/normal",        "F",                "Latidos híbridos — marcador de riesgo"),
    ("Q", "No clasificable / artefactos",     "Q, ?, /, f, !",    "Ruido, marcapasos, artefactos"),
]
for i, row_data in enumerate(aami):
    row = t2.rows[i + 1]
    _shade(row, "D6E4F0" if i % 2 == 0 else "FFFFFF")
    for ci, v in enumerate(row_data):
        _cell(row.cells[ci], v, bold=(ci == 0), center=(ci == 0))

blank(doc)

# ─────────────────────────────────────────────────────────────────────────────
h3(doc, "3.1.3  Fase 1: Adquisición y preprocesamiento de señales ECG")
# ─────────────────────────────────────────────────────────────────────────────

h4(doc, "Diseño del protocolo de preprocesamiento")

txt(doc,
    "Una de las decisiones de diseño más importantes del proyecto fue la elección "
    "de la longitud de la ventana de segmentación. En trabajos previos con esta base "
    "de datos, la longitud estándar empleada suele ser de 71 a 90 muestras (~200 ms). "
    "Sin embargo, la onda P del ECG —cuya morfología es determinante para distinguir "
    "los latidos supraventriculares ectópicos (clase S) de los normales— aparece "
    "típicamente entre 100 y 120 ms antes del pico R. Con una ventana de 71 muestras "
    "(35 muestras antes del pico a 360 Hz equivalen a 97 ms), la onda P quedaba "
    "parcialmente fuera de la ventana de análisis, limitando la capacidad del modelo "
    "para aprender las características morfológicas de la clase S.")

txt(doc,
    "Se diseñó un protocolo de preprocesamiento con ventana ampliada a 128 muestras, "
    "con 50 muestras anteriores al pico R (139 ms) y 78 posteriores (217 ms), "
    "capturando la onda P de forma completa. Esta ventana de 355 ms también captura "
    "el complejo QRS y la onda T para la mayoría de frecuencias cardíacas clínicas "
    "(40–150 lpm).")

blank(doc)
h4(doc, "Pipeline de preprocesamiento")

txt(doc, "El pipeline consta de cinco etapas aplicadas secuencialmente a cada registro:", sa=3)

num_list(doc, [
    ("Filtrado pasa-banda Butterworth: ",
     "Filtro IIR de orden 4 con banda de paso 0.5–40 Hz. El límite inferior "
     "elimina la deriva de línea base (baseline wander) causada por la respiración. "
     "El límite superior elimina el ruido eléctrico de 50/60 Hz. Se empleó "
     "filtrado bidireccional (filtfilt) para garantizar fase cero y evitar "
     "desplazamiento temporal de los complejos QRS."),
    ("Detección de picos R: ",
     "En el modo MIT-BIH se utilizaron directamente las anotaciones de los "
     "cardiólogos (ann.sample) como posiciones de los picos R. En el modo CSV "
     "(señales externas) se emplea el detector de picos locales de SciPy con "
     "umbral adaptativo y distancia mínima equivalente a 200 lpm."),
    ("Segmentación de latidos: ",
     "Para cada pico R en la posición k se extrajo la subsecuencia "
     "x[k − 50 : k + 78], obteniendo un vector de 128 muestras. Los latidos "
     "cuya ventana excedía los límites del registro fueron descartados."),
    ("Normalización z-score: ",
     "Cada segmento b fue normalizado individualmente: "
     "b̂ᵢ = (bᵢ − μb) / σb, donde μb y σb son la media y la desviación estándar "
     "del propio segmento. Esta normalización elimina las variaciones de amplitud "
     "entre pacientes y electrodos. Los segmentos con σb = 0 (señal plana) "
     "fueron descartados."),
    ("Mapeo de etiquetas AAMI EC57: ",
     "La etiqueta original de cada anotación fue mapeada a las cinco clases "
     "AAMI EC57 mediante el diccionario de correspondencia de la Tabla 3.2. "
     "Las anotaciones sin correspondencia fueron descartadas."),
])

blank(doc)
txt(doc,
    "El dataset resultante fue almacenado en formato NumPy: X_128.npy con dimensiones "
    "94,618 × 128 y y_128.npy con 94,618 etiquetas enteras de 0 a 4.")

blank(doc)
cap_tab(doc, "Tabla 3.3. Distribución de clases en el dataset procesado (MIT-BIH, estándar AAMI EC57).")
t3 = new_table(doc, 7, 4)
hdr(t3.rows[0], ["Clase", "Descripción", "Latidos", "Porcentaje"])
dist = [
    ("N", "Normal",                    "75,040", "79.3 %"),
    ("V", "Ventricular ectópico",       "7,235",  "7.6 %"),
    ("Q", "No clasificable",            "8,515",  "9.0 %"),
    ("S", "Supraventricular ectópico",  "3,026",  "3.2 %"),
    ("F", "Fusión ventricular/normal",    "802",  "0.8 %"),
]
for i, row_data in enumerate(dist):
    row = t3.rows[i + 1]
    _shade(row, "D6E4F0" if i % 2 == 0 else "FFFFFF")
    for ci, v in enumerate(row_data):
        _cell(row.cells[ci], v, bold=(ci == 0), center=(ci in (0, 2, 3)))
_shade(t3.rows[6], "BDD7EE")
for ci, v in enumerate(["Total", "—", "94,618", "100.0 %"]):
    _cell(t3.rows[6].cells[ci], v, bold=True, center=(ci in (0, 2, 3)))

blank(doc)
txt(doc,
    "El marcado desbalance de clases (clase N: 79.3 %, clase F: 0.8 %) fue una "
    "consideración crítica en el diseño del entrenamiento. Para mitigarlo se "
    "emplearon Focal Loss con pesos de clase y WeightedRandomSampler, "
    "descritos en la sección 3.1.5.")

blank(doc)

# ─────────────────────────────────────────────────────────────────────────────
h3(doc, "3.1.4  Fase 2: Diseño y entrenamiento del modelo neuronal base")
# ─────────────────────────────────────────────────────────────────────────────

h4(doc, "Arquitectura ECG_ResNet_SE")

txt(doc,
    "Se diseñó una red neuronal convolucional unidimensional con conexiones residuales "
    "y mecanismo de atención por canal, denominada ECG_ResNet_SE. La arquitectura "
    "combina bloques residuales (ResBlock1D) [4] con un bloque Squeeze-and-Excitation "
    "(SEBlock1D) [5], adaptados a señales temporales unidimensionales.")

txt(doc,
    "La selección de esta arquitectura sobre una CNN convencional responde a dos "
    "limitaciones: (1) las redes puramente convolucionales profundas sufren del "
    "desvanecimiento del gradiente, dificultando el aprendizaje de características "
    "a múltiples escalas; (2) todos los canales de características contribuyen por "
    "igual al vector de representación final. Las conexiones residuales resuelven "
    "el primer problema y el bloque SE resuelve el segundo.")

blank(doc)
cap_tab(doc, "Tabla 3.4. Arquitectura de la red neuronal ECG_ResNet_SE. "
        "Las dimensiones indican (canales, longitud temporal).")
t4 = new_table(doc, 11, 4)
hdr(t4.rows[0], ["Bloque", "Operación", "Configuración", "Salida"])
arch = [
    ("Entrada",   "—",                                    "—",                                    "(1, 128)"),
    ("Stem",      "Conv1d + BatchNorm + ReLU",             "32 filtros, kernel=7, padding=3",      "(32, 128)"),
    ("Layer 1",   "ResBlock1D(32→32) + MaxPool(2)",        "kernel=5, shortcut: identidad",        "(32, 64)"),
    ("Layer 2",   "ResBlock1D(32→64) + MaxPool(2)",        "kernel=5, shortcut: Conv 1×1",         "(64, 32)"),
    ("Layer 3",   "ResBlock1D(64→128) + MaxPool(2)",       "kernel=5, shortcut: Conv 1×1",         "(128, 16)"),
    ("SE Block",  "SEBlock1D(128, ratio=8)",               "FC 128→16→128, Sigmoid",               "(128, 16)"),
    ("Pool",      "AdaptiveAvgPool1d(1) + Flatten",        "—",                                    "128"),
    ("FC1",       "Linear + ReLU + Dropout(0.4)",          "128→64",                               "64"),
    ("Salida",    "Linear",                                "64→5",                                 "5 logits"),
    ("Total",     "188,469 parámetros entrenables",        "—",                                    "—"),
]
for i, row_data in enumerate(arch):
    row = t4.rows[i + 1]
    bg = "BDD7EE" if i == 9 else ("D6E4F0" if i % 2 == 0 else "FFFFFF")
    _shade(row, bg)
    for ci, v in enumerate(row_data):
        _cell(row.cells[ci], v, bold=(ci == 0 or i == 9), center=(ci in (2, 3)))

blank(doc)

h4(doc, "Bloque ResBlock1D")

txt(doc,
    "Cada bloque residual aplica dos convoluciones con Batch Normalization y ReLU, "
    "sumando la entrada original al resultado mediante una conexión de salto "
    "(shortcut connection). La salida del bloque es:")

txt(doc, "       y = ReLU( F(x) + S(x) )", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

txt(doc,
    "donde F(x) es la transformación de las dos convoluciones y S(x) es el shortcut: "
    "una proyección Conv 1×1 + BN cuando el número de canales cambia, o la identidad "
    "en caso contrario. La suma residual garantiza que el gradiente fluya directamente "
    "hacia capas anteriores, evitando el desvanecimiento en redes profundas.")

blank(doc)
h4(doc, "Bloque SEBlock1D (Squeeze-and-Excitation)")

txt(doc,
    "El bloque SE aprende un vector de importancia por canal de forma dinámica. "
    "Dado el mapa de características X con C canales y longitud temporal L, "
    "opera en tres etapas:")

num_list(doc, [
    ("Squeeze: ",
     "Se comprime la dimensión temporal mediante promedio global, generando "
     "un descriptor por canal: zc = (1/L) Σ Xc,l para l = 1, ..., L."),
    ("Excitation: ",
     "El descriptor pasa por dos capas FC que aprenden relaciones entre canales: "
     "s = σ(W₂ · δ(W₁ · z)), donde W₁ ∈ R^(C/r × C), W₂ ∈ R^(C × C/r), "
     "δ es ReLU, σ es sigmoide y r=8 es el factor de reducción."),
    ("Recalibración: ",
     "Cada canal del mapa de características es multiplicado por su escalar "
     "de importancia: X̃c,l = sc · Xc,l. Esto permite que la red enfatice "
     "los canales más relevantes para cada tipo de morfología de latido."),
])

blank(doc)

# ─────────────────────────────────────────────────────────────────────────────
h3(doc, "3.1.5  Estrategia de entrenamiento del modelo base")
# ─────────────────────────────────────────────────────────────────────────────

txt(doc,
    "El dataset fue dividido en tres particiones con estratificación por clase: "
    "70 % para entrenamiento (66,233 latidos), 15 % para validación (14,192 latidos) "
    "y 15 % como conjunto de prueba independiente (14,193 latidos). El conjunto de "
    "prueba no fue utilizado en ninguna decisión de diseño; se reservó exclusivamente "
    "para la evaluación final.")

blank(doc)
h4(doc, "Función de pérdida: Focal Loss")

txt(doc,
    "El desbalance de clases hace que la entropía cruzada estándar tienda a ignorar "
    "las clases minoritarias. Se empleó la Focal Loss [6]:")

txt(doc,
    "       FL(pt) = −αt · (1 − pt)^γ · log(pt)",
    bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

txt(doc,
    "donde pt es la probabilidad predicha para la clase correcta y γ = 2 reduce la "
    "contribución al gradiente de los ejemplos fáciles, concentrando el aprendizaje "
    "en los difíciles. Los pesos αc fueron calculados como la raíz cuadrada del "
    "inverso de la frecuencia de cada clase, normalizados:")

txt(doc,
    "       αc = ( √(1/nc) / Σk √(1/nk) ) × 5",
    bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

txt(doc,
    "donde nc es el número de latidos de la clase c en el entrenamiento. "
    "La elección de la raíz cuadrada en lugar de la inversa directa (1/nc) fue "
    "resultado de experimentación: con 1/nc el peso de la clase F resultó 3.38 veces "
    "mayor que el de N, provocando sobrepredicción masiva de la clase F "
    "(precisión = 0.49, recall = 0.87). Con √(1/nc) los pesos son más moderados "
    "y el modelo logra un balance superior entre precisión y recall en todas las clases.")

blank(doc)
h4(doc, "Muestreo ponderado y aumentación de datos")

txt(doc,
    "Adicionalmente a la Focal Loss, se empleó un WeightedRandomSampler para "
    "asegurar que cada minibatch contenga una representación balanceada de todas "
    "las clases. Para mejorar la generalización, se aplicaron dos transformaciones "
    "estocásticas durante el entrenamiento únicamente:")

bullet(doc, [
    "Ruido gaussiano aditivo: b̃ᵢ = bᵢ + εᵢ, donde εᵢ ~ N(0, 0.01²). "
    "Simula el ruido de medición del electrodo.",
    "Escalado de amplitud: b̃ᵢ = s · bᵢ, donde s ~ U(0.9, 1.1). "
    "Simula variaciones en la impedancia electrodo-piel entre pacientes.",
])

blank(doc)
h4(doc, "Optimizador, programador de tasa de aprendizaje y parada temprana")

txt(doc,
    "Se utilizó el optimizador AdamW [7] con tasa de aprendizaje inicial "
    "α = 10⁻³ y decaimiento de peso λ = 10⁻⁴. Para el programador de tasa "
    "de aprendizaje se empleó OneCycleLR: durante el 10 % inicial de las "
    "iteraciones la tasa aumenta linealmente desde α/10 hasta α (fase de "
    "calentamiento), y a partir de ese punto decrece suavemente siguiendo "
    "una curva coseno hasta α/10⁴ al final del entrenamiento.")

txt(doc,
    "Se implementó parada temprana con paciencia de 8 épocas monitorizando el "
    "F1-Score macro en validación. El entrenamiento máximo se fijó en 50 épocas. "
    "El modelo resultante fue guardado como models/ecg_resnet_se_base.pth y "
    "constituye el modelo estático en la evaluación comparativa.")

blank(doc)
cap_tab(doc, "Tabla 3.5. Hiperparámetros del entrenamiento base del modelo ECG_ResNet_SE (Fase 2).")
t5 = new_table(doc, 9, 3)
hdr(t5.rows[0], ["Hiperparámetro", "Valor", "Justificación"])
hp = [
    ("Épocas máximas",        "50",             "Convergencia observada antes de la época 35 en todos los runs"),
    ("Parada temprana",       "Paciencia = 8 épocas (F1 macro val.)", "Previene sobreajuste sin interrumpir prematuramente"),
    ("Tamaño de minibatch",   "128",            "Balance entre estabilidad del gradiente y velocidad en GPU T4"),
    ("Optimizador",           "AdamW",          "Decaimiento de peso desacoplado; mejor generalización que Adam estándar"),
    ("Tasa de aprendizaje",   "α = 1×10⁻³",    "Valor estándar para AdamW en clasificación de señales biomédicas"),
    ("Decaimiento de peso",   "λ = 1×10⁻⁴",    "Regularización L2 implícita para prevenir sobreajuste"),
    ("Programador de LR",     "OneCycleLR (warmup 10 %, cosine decay)", "Convergencia más rápida y mejor F1 final que StepLR"),
    ("Función de pérdida",    "Focal Loss (γ=2, α=√(1/nc))", "Enfoca el gradiente en ejemplos difíciles de clases minoritarias"),
]
for i, row_data in enumerate(hp):
    row = t5.rows[i + 1]
    _shade(row, "D6E4F0" if i % 2 == 0 else "FFFFFF")
    _cell(row.cells[0], row_data[0], bold=True, center=False)
    _cell(row.cells[1], row_data[1], center=True)
    _cell(row.cells[2], row_data[2], center=False, sz=9)

blank(doc)

# ─────────────────────────────────────────────────────────────────────────────
h3(doc, "3.1.6  Fase 3: Módulo de reentrenamiento continuo con Replay Buffer")
# ─────────────────────────────────────────────────────────────────────────────

h4(doc, "El problema del Catastrophic Forgetting")

txt(doc,
    "Un modelo neuronal entrenado sobre un conjunto de datos inicial pierde "
    "gradualmente el conocimiento adquirido cuando se actualiza con nuevos datos "
    "provenientes de una distribución diferente. Este fenómeno, denominado "
    "Catastrophic Forgetting [8], ocurre porque el descenso de gradiente sobre "
    "los nuevos ejemplos sobreescribe los pesos que codificaban el conocimiento "
    "anterior. En el contexto clínico de este proyecto, esto equivale a que el "
    "modelo, al adaptarse a los patrones de un nuevo paciente, pierda la capacidad "
    "de clasificar correctamente los patrones de pacientes anteriores.")

blank(doc)
h4(doc, "Implementación del Replay Buffer")

txt(doc,
    "La solución implementada es la técnica de Experience Replay [9]: se almacena "
    "una muestra representativa de los datos anteriores en un buffer de memoria y "
    "se mezcla con los nuevos datos en cada iteración de reentrenamiento. El Replay "
    "Buffer implementado funciona como una cola FIFO de capacidad máxima M = 2,000 "
    "latidos. Cuando se agregan nuevos latidos y la capacidad es superada, se "
    "descartan los más antiguos. El buffer se persiste entre sesiones en el archivo "
    "models/replay_buffer.npz.")

blank(doc)
h4(doc, "Algoritmo de reentrenamiento incremental")

txt(doc,
    "En cada sesión de reentrenamiento sobre un nuevo conjunto de latidos "
    "(X_new, y_new), el algoritmo procede de la siguiente manera:")

num_list(doc, [
    ("Incorporación al buffer: ",
     "Se agregan los nuevos latidos al Replay Buffer."),
    ("Muestreo de experiencias pasadas: ",
     "Se extraen aleatoriamente n_r = ⌊|X_new| × 0.5⌋ latidos del buffer, "
     "obteniendo (X_rep, y_rep)."),
    ("Mezcla de datos: ",
     "Se construye el conjunto combinado: X_comb = [X_new ; X_rep], "
     "y_comb = [y_new ; y_rep], y se permuta aleatoriamente."),
    ("Actualización de pesos: ",
     "Se entrena el modelo durante 5 épocas con el optimizador Adam "
     "(α = 10⁻⁴, función de pérdida de entropía cruzada) y minibatches "
     "de 64 latidos. La tasa de aprendizaje reducida (10⁻⁴ frente a 10⁻³ "
     "del entrenamiento base) es fundamental para preservar las representaciones "
     "aprendidas: una tasa alta provocaría sobreescritura del conocimiento previo."),
])

txt(doc,
    "La proporción 50 % datos nuevos / 50 % buffer representa el equilibrio entre "
    "plasticidad (capacidad de aprender nuevos patrones) y estabilidad (retención "
    "del conocimiento previo), denominado stability-plasticity trade-off en la "
    "literatura de Continual Learning.")

blank(doc)
cap_tab(doc, "Tabla 3.6. Parámetros del módulo de reentrenamiento continuo con Replay Buffer.")
t6 = new_table(doc, 8, 3)
hdr(t6.rows[0], ["Parámetro", "Valor", "Justificación"])
rp = [
    ("Capacidad del Replay Buffer", "2,000 latidos",      "Representa la diversidad histórica sin exceder la RAM disponible"),
    ("Política de reemplazo",       "FIFO",               "Los latidos más recientes tienen prioridad sobre los más antiguos"),
    ("Proporción buffer / nuevos",  "50 % / 50 % / época","Balance entre estabilidad y plasticidad del modelo"),
    ("Tasa de aprendizaje",         "α = 1×10⁻⁴",        "Diez veces menor que el entrenamiento base; cambios pequeños y seguros"),
    ("Épocas por sesión",           "5",                  "Ajuste incremental sin sobreajuste a los datos de la sesión"),
    ("Tamaño de minibatch",         "64 latidos",         "Consistente con el entrenamiento base"),
    ("Lotes en evaluación formal",  "6 (secuenciales)",   "~15,770 latidos nuevos por lote; lote 0 solo inicializa el buffer"),
]
for i, row_data in enumerate(rp):
    row = t6.rows[i + 1]
    _shade(row, "D6E4F0" if i % 2 == 0 else "FFFFFF")
    _cell(row.cells[0], row_data[0], bold=True, center=False)
    _cell(row.cells[1], row_data[1], center=True)
    _cell(row.cells[2], row_data[2], center=False, sz=9)

blank(doc)
txt(doc,
    "Para la evaluación formal del módulo (Fase 3), el dataset completo de 94,618 "
    "latidos fue dividido en 6 lotes secuenciales de aproximadamente 15,770 latidos "
    "cada uno, representando la llegada progresiva de datos de distintos grupos de "
    "pacientes. El lote 0 únicamente inicializó el Replay Buffer sin reentrenar, "
    "actuando como referencia de partida. Los lotes 1 a 5 fueron procesados "
    "secuencialmente con el algoritmo incremental descrito.")

blank(doc)

# ─────────────────────────────────────────────────────────────────────────────
h3(doc, "3.1.7  Fase 4: Evaluación comparativa formal")
# ─────────────────────────────────────────────────────────────────────────────

h4(doc, "Protocolo de evaluación")

txt(doc,
    "La evaluación comparativa formal se realizó sobre el dataset completo "
    "(94,618 latidos), comparando el modelo estático (pesos fijos del entrenamiento "
    "base, sin ninguna actualización posterior) contra el modelo adaptativo (pesos "
    "del modelo tras el reentrenamiento continuo sobre los 6 lotes). Las métricas "
    "calculadas para cada modelo fueron exactitud (Accuracy), F1-Score macro, "
    "precisión y sensibilidad por clase AAMI EC57.")

blank(doc)
h4(doc, "Definición formal de métricas")

txt(doc, "Las métricas empleadas se definen formalmente como:", sa=3)

bullet(doc, [
    "Exactitud (Accuracy): fracción de latidos correctamente clasificados sobre el total.",
    "Precisión por clase c: Pc = TPc / (TPc + FPc). Fracción de predicciones de la clase c que son correctas.",
    "Sensibilidad por clase c: Rc = TPc / (TPc + FNc). Fracción de latidos reales de la clase c correctamente detectados.",
    "F1-Score por clase c: F1c = 2 · Pc · Rc / (Pc + Rc). Media armónica entre precisión y sensibilidad.",
    "F1-Score macro: promedio no ponderado del F1-Score sobre las cinco clases. Pondera por igual las clases minoritarias y mayoritarias.",
])

blank(doc)
h4(doc, "Análisis de significancia estadística: Prueba de McNemar")

txt(doc,
    "Para determinar si la diferencia de rendimiento entre ambos modelos es "
    "estadísticamente significativa, se aplicó la prueba de McNemar [10]. Esta "
    "prueba evalúa si dos clasificadores cometen errores distintos sobre los mismos "
    "ejemplos y es el método estándar para comparar clasificadores sobre un mismo "
    "conjunto de datos de prueba.")

txt(doc,
    "Si b es el número de latidos clasificados correctamente solo por el modelo "
    "adaptativo y c el número clasificado correctamente solo por el modelo estático, "
    "el estadístico de prueba es:")

txt(doc,
    "       χ²_McNemar = (|b − c| − 1)² / (b + c)",
    bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

txt(doc,
    "que sigue una distribución χ² con 1 grado de libertad bajo la hipótesis nula "
    "de que ambos clasificadores tienen el mismo rendimiento. Se rechaza la hipótesis "
    "nula si p < 0.05. Dado que la función scipy.stats.mcnemar no está disponible "
    "en versiones recientes de SciPy, el estadístico fue calculado manualmente.")

blank(doc)

# ─────────────────────────────────────────────────────────────────────────────
h3(doc, "3.1.8  Interfaz de usuario")
# ─────────────────────────────────────────────────────────────────────────────

txt(doc,
    "Se desarrolló una interfaz de usuario web utilizando el framework Streamlit "
    "(src/ui/app.py), accesible desde cualquier navegador al ejecutar el comando "
    "streamlit run src/ui/app.py desde el directorio raíz del proyecto. La "
    "interfaz integra el pipeline completo de análisis, inferencia y reentrenamiento "
    "continuo en una sola aplicación sin necesidad de conocimientos de programación.")

blank(doc)
txt(doc, "La interfaz se organiza en tres pestañas principales:", sa=3)

num_list(doc, [
    ("Análisis e Inferencia: ",
     "Permite seleccionar cualquiera de los 48 registros MIT-BIH o cargar una "
     "señal ECG propia en formato CSV. Ejecuta automáticamente el pipeline completo "
     "(filtrado, segmentación con ventana de 128 muestras, normalización e inferencia) "
     "y presenta una gráfica interactiva de la señal ECG con los picos R codificados "
     "por color según la clase predicha, la distribución de clases, las métricas por "
     "clase y una tabla detallada con la predicción y confianza para cada latido. "
     "Incluye el módulo de reentrenamiento con parámetros ajustables por el usuario "
     "(épocas, tasa de aprendizaje, tamaño del buffer)."),
    ("Historial de Reentrenamiento: ",
     "Presenta la evolución temporal del rendimiento del modelo adaptativo a lo "
     "largo de todas las sesiones interactivas, con gráficas de exactitud y F1-Score "
     "macro antes y después de cada sesión, tabla exportable en CSV y opción de "
     "reinicio del historial. El historial se persiste en models/retrain_history.json."),
    ("Análisis Comparativo: ",
     "Muestra la comparación formal entre el modelo estático y el modelo adaptativo. "
     "Los KPIs del modelo adaptativo se calculan dinámicamente como el promedio de "
     "los últimos resultados por registro en el historial. En ausencia de historial, "
     "se muestran los resultados de la evaluación formal de la Fase 3 como referencia. "
     "Incluye gráficas por registro de F1-Score y exactitud finales, así como la "
     "evolución del F1-Score promedio acumulado a lo largo de las sesiones."),
])

blank(doc)

# ─────────────────────────────────────────────────────────────────────────────
# NOTA FINAL
# ─────────────────────────────────────────────────────────────────────────────
txt(doc,
    "Nota: La sección 3.2 (Resultados) será completada una vez finalizado el "
    "proceso de reentrenamiento con los datos actualizados.",
    italic=True, size=11)

# ─────────────────────────────────────────────────────────────────────────────
# GUARDAR
# ─────────────────────────────────────────────────────────────────────────────
doc.save(OUT)
print(f"\nDocumento generado en:\n  {OUT}\n")
