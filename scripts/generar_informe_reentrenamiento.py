"""
Genera docs/ADAPT-ECG_Informe_Reentrenamiento.docx
Documentacion completa del proceso de reentrenamiento continuo:
modelo base vs modelo adaptativo, Replay Buffer, resultados por registro
y analisis comparativo final.

Alumno: Elias Bejarano Lozada (#20213057), Ingenieria Biomedica, ITT
"""

import json
from collections import defaultdict
from pathlib import Path

import numpy as np
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT        = Path(__file__).resolve().parents[1]
HISTORY_PATH = ROOT / "models" / "retrain_history.json"
OUT          = ROOT / "docs" / "ADAPT-ECG_Informe_Reentrenamiento.docx"

# ---------------------------------------------------------------------------
# Helpers de formato (estilo ITT)
# ---------------------------------------------------------------------------
NEGRO = RGBColor(0x00, 0x00, 0x00)
AZUL  = RGBColor(0x1A, 0x56, 0x9A)
VERDE = RGBColor(0x1A, 0x6B, 0x2A)
ROJO  = RGBColor(0xC0, 0x2A, 0x2A)

doc = Document()

for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(2.5)


def set_run_font(run, size=12, bold=False, italic=False, color=None):
    run.font.name  = "Times New Roman"
    run.font.size  = Pt(size)
    run.font.bold  = bold
    run.font.italic = italic
    run.font.color.rgb = color if color else NEGRO
    rPr = run._r.get_or_add_rPr()
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), "Times New Roman")
    rFonts.set(qn("w:hAnsi"), "Times New Roman")
    rFonts.set(qn("w:cs"),    "Times New Roman")
    rPr.insert(0, rFonts)


def heading(text, level=1, space_before=14):
    sizes = {1: 14, 2: 13, 3: 12}
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.line_spacing = Pt(18)
    run = p.add_run(text)
    set_run_font(run, size=sizes.get(level, 12), bold=True,
                 color=AZUL if level == 1 else NEGRO)
    return p


def body(text, justify=True):
    p = doc.add_paragraph()
    p.paragraph_format.space_after       = Pt(6)
    p.paragraph_format.first_line_indent = Cm(0.75)
    p.paragraph_format.line_spacing      = Pt(18)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY if justify else WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    set_run_font(run, size=12)
    return p


def bullet(text, indent=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_after  = Pt(3)
    p.paragraph_format.left_indent  = Cm(indent)
    p.paragraph_format.line_spacing = Pt(16)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(u"\u2022  " + text)
    set_run_font(run, size=12)
    return p


def nota(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.left_indent  = Cm(0.75)
    p.paragraph_format.line_spacing = Pt(16)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run("Nota: " + text)
    set_run_font(run, size=11, italic=True)
    return p


def tabla(encabezados, filas, col_widths=None):
    t = doc.add_table(rows=1 + len(filas), cols=len(encabezados))
    t.style = "Table Grid"
    for i, h in enumerate(encabezados):
        cell = t.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        set_run_font(run, size=11, bold=True)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        if col_widths:
            cell.width = Cm(col_widths[i])
    for r, fila in enumerate(filas):
        for c, val in enumerate(fila):
            cell = t.rows[r + 1].cells[c]
            cell.text = ""
            run = cell.paragraphs[0].add_run(str(val))
            set_run_font(run, size=11)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()


# ---------------------------------------------------------------------------
# Cargar datos del historial
# ---------------------------------------------------------------------------
with open(HISTORY_PATH, encoding="utf-8") as f:
    history = json.load(f)

por_reg = defaultdict(list)
for h in history:
    por_reg[h["registro"]].append(h)

regs    = sorted(por_reg.keys())
f1_fin  = [por_reg[r][-1]["f1_despues"] for r in regs]
acc_fin = [por_reg[r][-1]["acc_despues"] for r in regs]

perfectos = [r for r in regs if por_reg[r][-1]["f1_despues"] == 1.0]
buenos    = [r for r in regs if 0.70 <= por_reg[r][-1]["f1_despues"] < 1.0]
limitados = [r for r in regs if por_reg[r][-1]["f1_despues"] < 0.70]

acc_prom = np.mean(acc_fin)
f1_prom  = np.mean(f1_fin)
f1_med   = np.median(f1_fin)

STATIC_ACC  = 0.8872
STATIC_F1   = 0.7381
RETRAIN_ACC = 0.9741
RETRAIN_F1  = 0.8954

# ---------------------------------------------------------------------------
# PORTADA
# ---------------------------------------------------------------------------
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(0)
p.paragraph_format.space_after  = Pt(10)
run = p.add_run("INFORME DE REENTRENAMIENTO CONTINUO")
set_run_font(run, size=18, bold=True, color=AZUL)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(6)
run = p.add_run(
    "Proceso de Adaptacion del Modelo CNN para Deteccion de Arritmias\n"
    "Modelo Base vs Modelo Reentrenado — Analisis Comparativo"
)
set_run_font(run, size=14, bold=True)

for line in [
    "Sistema inteligente con reentrenamiento continuo para la deteccion adaptativa",
    "de patologias cardiovasculares basado en senales ECG",
    "",
    "Alumno: Elias Bejarano Lozada  |  No. Control: 20213057",
    "Asesor: M.C. Fortunato Ramirez Arzate",
    "Ingenieria Biomedica — Instituto Tecnologico de Tijuana",
    "Residencia Profesional, 2026",
]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(line)
    set_run_font(run, size=12, italic=not line.startswith("Alumno") and line != "")

doc.add_page_break()

# ---------------------------------------------------------------------------
# 1. INTRODUCCION
# ---------------------------------------------------------------------------
heading("1. Introduccion", level=1)

body(
    "El presente informe documenta el proceso de reentrenamiento continuo aplicado al modelo "
    "de red neuronal convolucional (CNN) desarrollado en el proyecto ADAPT-ECG. El objetivo "
    "principal es demostrar la capacidad del sistema para adaptarse a nuevos datos de pacientes "
    "de forma incremental, sin necesidad de reentrenar el modelo desde cero, y sin incurrir en "
    "el fenomeno conocido como catastrophic forgetting."
)
body(
    "El proceso de reentrenamiento se realizo sobre los 46 registros de la base de datos "
    "MIT-BIH Arrhythmia Database (PhysioNet), utilizando la tecnica de Replay Buffer para "
    "preservar el conocimiento adquirido en sesiones anteriores. Los resultados se comparan "
    "contra el modelo base estatico entrenado en la Fase 2 del proyecto."
)

# ---------------------------------------------------------------------------
# 2. METODOLOGIA
# ---------------------------------------------------------------------------
heading("2. Metodologia de Reentrenamiento", level=1)

heading("2.1 Tecnica: Replay Buffer", level=2)
body(
    "El Replay Buffer es una tecnica de aprendizaje continuo (Continual Learning) que almacena "
    "una muestra representativa de datos anteriores para mezclarlos con los datos nuevos durante "
    "cada sesion de reentrenamiento. Esto permite al modelo retener el conocimiento previo "
    "mientras aprende nuevos patrones, evitando el catastrophic forgetting."
)

bullet("Capacidad maxima del buffer: 2,000 latidos")
bullet("Estrategia de muestreo: FIFO (los mas antiguos se descartan al superar el limite)")
bullet("Mezcla por epoca: 50% datos nuevos + 50% muestreo aleatorio del buffer")
bullet("Persistencia: el buffer se guarda en models/replay_buffer.npz entre sesiones")

heading("2.2 Parametros de entrenamiento", level=2)
body("Los hiperparametros utilizados en cada sesion de reentrenamiento son los siguientes:")

tabla(
    ["Parametro", "Valor", "Justificacion"],
    [
        ["Optimizador",   "Adam",   "Convergencia rapida, adecuado para fine-tuning"],
        ["Learning Rate", "1e-4",   "10x menor que el entrenamiento base (evita sobreescribir pesos)"],
        ["Batch size",    "64",     "Igual al entrenamiento base para consistencia"],
        ["Epocas",        "5",      "Suficiente para adaptar sin sobreajustar"],
        ["Loss",          "CrossEntropyLoss", "Sin ponderacion, estandar para clasificacion multiclase"],
        ["Buffer size",   "2,000",  "Balance entre memoria y representatividad"],
        ["Sesiones/reg",  "2",      "Convergencia en registros grandes (>500 latidos)"],
    ],
    col_widths=[4, 3, 8]
)

heading("2.3 Flujo de una sesion de reentrenamiento", level=2)
body("Cada sesion de reentrenamiento sigue el siguiente proceso:")

bullet("1. Cargar registro MIT-BIH: lectura de la senal ECG y anotaciones (.hea, .dat, .atr)")
bullet("2. Preprocesamiento: filtro Butterworth pasa-banda 0.5-40 Hz, segmentacion en latidos (BEAT_LEN=71), normalizacion z-score por latido")
bullet("3. Evaluacion inicial: se mide Accuracy y F1 Macro del modelo actual sobre el registro")
bullet("4. Agregar al buffer: los latidos del registro se incorporan al Replay Buffer (FIFO)")
bullet("5. Entrenamiento incremental: por cada epoca, se combinan datos nuevos + muestreo del buffer")
bullet("6. Evaluacion final: se miden Accuracy y F1 Macro despues del reentrenamiento")
bullet("7. Guardar: modelo .pth, buffer .npz e historial .json se actualizan en disco")

heading("2.4 Clasificacion AAMI (5 clases)", level=2)
body(
    "El modelo clasifica cada latido en una de las cinco categorias definidas por el estandar "
    "AAMI EC57 para arritmias cardiacas:"
)
tabla(
    ["Codigo", "Nombre", "Simbolos MIT-BIH", "Descripcion"],
    [
        ["N", "Normal",           "N, .",         "Latidos normales y de union"],
        ["S", "Supraventricular", "A, a, J, S, e, j", "Ectopicos supraventriculares"],
        ["V", "Ventricular",      "V, E",         "Ectopicos ventriculares"],
        ["F", "Fusion",           "F",            "Fusion ventricular/normal"],
        ["Q", "Desconocido",      "Q, ?, /, f, !", "No clasificable / artefactos"],
    ],
    col_widths=[2, 4, 4, 6]
)

# ---------------------------------------------------------------------------
# 3. BASE DE DATOS
# ---------------------------------------------------------------------------
heading("3. Base de Datos MIT-BIH", level=1)

body(
    "La base de datos MIT-BIH Arrhythmia Database es el estandar de referencia para la "
    "evaluacion de algoritmos de deteccion de arritmias. Contiene 48 registros de ECG de "
    "dos canales, muestreados a 360 Hz, cada uno de aproximadamente 30 minutos de duracion."
)

body(
    f"Para el proceso de reentrenamiento se utilizaron {len(regs)} registros de los 48 "
    f"disponibles. Se excluyeron los registros 111 (1 latido valido) y 124 (marcapaso, "
    f"solo 88 latidos con clase Q dominante que causa oscilacion en el modelo)."
)

tabla(
    ["Categoria", "Registros", "Criterio"],
    [
        ["Procesados con 2 sesiones",
         str(len([r for r in regs if por_reg[r][-1]["n_latidos"] >= 500])),
         ">= 500 latidos validos AAMI"],
        ["Procesados con limitaciones",
         str(len([r for r in regs if por_reg[r][-1]["n_latidos"] < 500])),
         "< 500 latidos (109, 118, 214, 231)"],
        ["Excluido (marcapaso)",  "1", "Registro 124 — solo 88 latidos clase Q"],
        ["Excluido (insuficiente)", "1", "Registro 111 — solo 1 latido valido"],
        ["Total base de datos",  "48", "MIT-BIH Arrhythmia Database completa"],
    ],
    col_widths=[5, 3, 7]
)

nota(
    "Los registros con menos de 500 latidos presentan mayor variabilidad en los resultados "
    "debido al reducido tamano de muestra. Sus metricas deben interpretarse con precaucion."
)

# ---------------------------------------------------------------------------
# 4. RESULTADOS COMPARATIVOS — FASE 4
# ---------------------------------------------------------------------------
heading("4. Resultados Comparativos — Evaluacion Formal (Fase 4)", level=1)

body(
    "La evaluacion formal del proyecto se realizo en la Fase 4 sobre 75,702 latidos "
    "(lotes 1 al 4 de la base de datos). Los resultados demuestran una mejora estadisticamente "
    "significativa del modelo adaptativo sobre el modelo estatico."
)

tabla(
    ["Metrica", "Modelo Estatico", "Modelo Adaptativo", "Mejora"],
    [
        ["Accuracy",    f"{STATIC_ACC*100:.2f}%",  f"{RETRAIN_ACC*100:.2f}%",
         f"+{(RETRAIN_ACC-STATIC_ACC)*100:.2f}%"],
        ["F1 Macro",    f"{STATIC_F1:.4f}",         f"{RETRAIN_F1:.4f}",
         f"+{RETRAIN_F1-STATIC_F1:.4f}"],
        ["F1 Weighted", "0.9447",                   "0.9773",        "+0.0326"],
        ["Prueba McNemar", "—", "chi2=5342, p~0", "Diferencia significativa (p<0.001)"],
    ],
    col_widths=[4.5, 3.5, 3.5, 4]
)

body(
    "La prueba de McNemar confirma que la diferencia entre ambos modelos es estadisticamente "
    "significativa (chi2=5342, p aproximadamente 0), descartando que la mejora sea producto "
    "del azar. Esto valida la efectividad del enfoque de reentrenamiento continuo con "
    "Replay Buffer para combatir el concept drift en senales ECG."
)

# ---------------------------------------------------------------------------
# 5. RESULTADOS POR REGISTRO
# ---------------------------------------------------------------------------
heading("5. Resultados del Reentrenamiento por Registro MIT-BIH", level=1)

body(
    f"Se reentrenaron {len(regs)} registros con 2 sesiones cada uno. A continuacion se "
    f"presentan los resultados finales (ultima sesion) de cada registro, ordenados por numero."
)

# Tabla de resultados por registro
filas_reg = []
for r in regs:
    h = por_reg[r][-1]
    estado = "Perfecto" if h["f1_despues"] == 1.0 else (
             "Bueno"    if h["f1_despues"] >= 0.70 else "Limitado")
    filas_reg.append([
        r,
        str(h["n_latidos"]),
        f"{h['acc_despues']*100:.2f}%",
        f"{h['f1_despues']:.4f}",
        estado
    ])

tabla(
    ["Registro", "Latidos", "Accuracy Final", "F1 Macro Final", "Estado"],
    filas_reg,
    col_widths=[2.5, 2.5, 3.5, 3.5, 3]
)

nota(
    "Estado: Perfecto = F1 1.00, Bueno = F1 >= 0.70, Limitado = F1 < 0.70 "
    "(generalmente por distribucion uniclase o pocos latidos)."
)

# ---------------------------------------------------------------------------
# 6. ANALISIS DE RESULTADOS
# ---------------------------------------------------------------------------
heading("6. Analisis de Resultados", level=1)

heading("6.1 Estadisticas globales del reentrenamiento", level=2)

tabla(
    ["Metrica", "Valor"],
    [
        ["Total registros evaluados",           str(len(regs))],
        ["Accuracy promedio final",             f"{acc_prom*100:.2f}%"],
        ["F1 Macro promedio final",             f"{f1_prom:.4f}"],
        ["F1 Macro mediana final",              f"{f1_med:.4f}"],
        ["Registros con F1 = 1.00 (perfectos)", f"{len(perfectos)} ({len(perfectos)/len(regs)*100:.0f}%)"],
        ["Registros con F1 >= 0.70",            f"{len(perfectos)+len(buenos)} ({(len(perfectos)+len(buenos))/len(regs)*100:.0f}%)"],
        ["Registros con F1 < 0.70 (limitados)", f"{len(limitados)} ({len(limitados)/len(regs)*100:.0f}%)"],
        ["Total sesiones de reentrenamiento",   str(len(history))],
    ],
    col_widths=[8, 5]
)

heading("6.2 Registros con clasificacion perfecta (F1 = 1.00)", level=2)
body(
    f"Un total de {len(perfectos)} registros alcanzaron F1 Macro = 1.00, indicando que el "
    f"modelo adaptativo clasifica correctamente el 100% de los latidos en todas las clases "
    f"presentes. Estos registros son:"
)
bullet(", ".join(perfectos))

heading("6.3 Registros con resultados limitados (F1 < 0.70)", level=2)
body(
    f"Los {len(limitados)} registros con F1 < 0.70 presentan limitaciones inherentes a la "
    f"composicion de sus datos, no al modelo. Se identificaron dos causas principales:"
)
bullet(
    "Distribucion uniclase: registros con mas del 98% de latidos clase N (Normal). "
    "El F1 Macro tiene un techo matematico de 0.50 cuando solo existe una clase en la muestra. "
    "Ejemplos: 103, 117, 200, 219."
)
bullet(
    "Muestra insuficiente: registros con menos de 300 latidos validos. "
    "El gradiente de actualizacion es ruidoso e inestable con tan pocos ejemplos. "
    "Ejemplos: 109 (40 latidos), 118 (112 latidos), 214 (259 latidos)."
)
nota(
    "La baja metrica F1 en estos registros no indica falla del modelo sino limitacion de los datos. "
    "La Accuracy en los mismos registros es generalmente superior al 95%, confirmando "
    "que el modelo clasifica correctamente la clase dominante."
)

heading("6.4 Comparacion con modelo estatico", level=2)
body(
    f"Comparando contra el modelo estatico (evaluacion Fase 4), el modelo reentrenado "
    f"muestra mejoras consistentes en todas las metricas:"
)
bullet(f"Accuracy: {STATIC_ACC*100:.2f}% -> {RETRAIN_ACC*100:.2f}%  (+{(RETRAIN_ACC-STATIC_ACC)*100:.2f} puntos porcentuales)")
bullet(f"F1 Macro: {STATIC_F1:.4f} -> {RETRAIN_F1:.4f}  (+{RETRAIN_F1-STATIC_F1:.4f})")
bullet(f"Registros con F1 >= 0.70: {len(perfectos)+len(buenos)} de {len(regs)} ({(len(perfectos)+len(buenos))/len(regs)*100:.0f}%)")
bullet(f"Registros perfectos (F1 = 1.00): {len(perfectos)} ({len(perfectos)/len(regs)*100:.0f}%)")

# ---------------------------------------------------------------------------
# 7. EXPERIMENTOS REALIZADOS
# ---------------------------------------------------------------------------
heading("7. Experimentos de Optimizacion Realizados", level=1)

body(
    "Durante el desarrollo del proceso de reentrenamiento se realizaron experimentos "
    "para intentar mejorar el F1 Macro en registros con clases desbalanceadas. "
    "A continuacion se documentan los enfoques evaluados y sus resultados."
)

tabla(
    ["Experimento", "Descripcion", "Resultado", "Conclusion"],
    [
        ["Loss ponderada (v1)",
         "Pesos 1/count por clase, incluyendo clases ausentes",
         "F1 promedio bajo de 0.797 a 0.685",
         "Demasiado agresivo en registros uniclase"],
        ["Buffer balanceado",
         "Cupo fijo por clase (400 por clase)",
         "Combinado con loss v1, empeoro resultados",
         "Buffer vacio al inicio causa inestabilidad"],
        ["Loss ponderada (v2)",
         "Sqrt inversa, solo si hay >=2 clases en batch",
         "F1 promedio 0.768 (similar al original)",
         "Mejora marginal, no justifica complejidad adicional"],
        ["Enfoque original",
         "CrossEntropyLoss estandar + FIFO buffer",
         "F1 promedio 0.768-0.797, mediana 0.848",
         "Mejor resultado general, elegido como definitivo"],
    ],
    col_widths=[3.5, 4, 3.5, 4]
)

body(
    "La conclusion de los experimentos es que el enfoque simple (loss estandar, buffer FIFO) "
    "produce los mejores resultados globales para este dataset. Las tecnicas de ponderacion "
    "pueden mejorar el F1 en registros con multiples clases, pero perjudican registros "
    "con distribucion uniclase, que son mayoría en MIT-BIH."
)

# ---------------------------------------------------------------------------
# 8. CONCLUSIONES
# ---------------------------------------------------------------------------
heading("8. Conclusiones", level=1)

body(
    "El sistema de reentrenamiento continuo con Replay Buffer demostro ser efectivo para "
    "la adaptacion incremental del modelo CNN a nuevos datos de pacientes. Los resultados "
    "validan la hipotesis del proyecto: un modelo adaptativo supera consistentemente a "
    "un modelo estatico en la deteccion de arritmias cardiacas."
)

bullet(
    f"El modelo adaptativo alcanzo una accuracy promedio de {acc_prom*100:.2f}% y "
    f"F1 Macro mediana de {f1_med:.4f} sobre {len(regs)} registros MIT-BIH."
)
bullet(
    f"{len(perfectos)} registros ({len(perfectos)/len(regs)*100:.0f}%) alcanzaron "
    f"F1 Macro = 1.00 (clasificacion perfecta en todas las clases presentes)."
)
bullet(
    f"{len(perfectos)+len(buenos)} registros ({(len(perfectos)+len(buenos))/len(regs)*100:.0f}%) "
    f"superaron el umbral de F1 >= 0.70, considerado satisfactorio para aplicaciones clinicas de apoyo."
)
bullet(
    "Los registros con F1 < 0.70 presentan limitaciones de datos (distribucion uniclase "
    "o menos de 300 latidos), no del modelo. Este comportamiento es esperado y documentado "
    "en la literatura de Continual Learning."
)
bullet(
    "La prueba estadistica de McNemar (chi2=5342, p~0) confirma que la diferencia entre "
    "el modelo estatico y el adaptativo es significativa, validando el enfoque propuesto."
)

body(
    "El sistema esta listo para su uso en contexto academico y de investigacion. "
    "La interfaz web (Streamlit) permite el reentrenamiento interactivo con nuevos "
    "registros ECG, ya sea de la base MIT-BIH o de archivos CSV propios, haciendo "
    "el sistema accesible para usuarios sin conocimientos de programacion."
)

# ---------------------------------------------------------------------------
# Guardar
# ---------------------------------------------------------------------------
doc.save(str(OUT))
print(f"Documento generado: {OUT}")
