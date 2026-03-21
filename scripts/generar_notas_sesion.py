"""
Genera documento Word con las notas de la sesión de preguntas sobre ADAPT-ECG.
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pathlib import Path

doc = Document()

# ── Estilos generales ─────────────────────────────────────────────────────────
style_normal = doc.styles["Normal"]
style_normal.font.name = "Calibri"
style_normal.font.size = Pt(11)

def titulo_doc(texto):
    p = doc.add_heading(texto, level=0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.runs[0]
    run.font.color.rgb = RGBColor(0x1A, 0x23, 0x7E)
    run.font.size = Pt(18)

def subtitulo(texto):
    p = doc.add_heading(texto, level=1)
    run = p.runs[0]
    run.font.color.rgb = RGBColor(0x1A, 0x23, 0x7E)
    run.font.size = Pt(14)

def subseccion(texto):
    p = doc.add_heading(texto, level=2)
    run = p.runs[0]
    run.font.color.rgb = RGBColor(0x15, 0x65, 0xC0)
    run.font.size = Pt(12)

def parrafo(texto):
    p = doc.add_paragraph(texto)
    p.paragraph_format.space_after = Pt(6)

def bullet(texto):
    p = doc.add_paragraph(texto, style="List Bullet")
    p.paragraph_format.space_after = Pt(3)

def codigo(texto):
    p = doc.add_paragraph()
    run = p.add_run(texto)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1B, 0x5E, 0x20)
    p.paragraph_format.left_indent = Inches(0.4)
    p.paragraph_format.space_after = Pt(4)
    shading = OxmlElement("w:shd")
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:color"), "auto")
    shading.set(qn("w:fill"), "F1F8E9")
    p._p.get_or_add_pPr().append(shading)

def tabla(encabezados, filas):
    t = doc.add_table(rows=1, cols=len(encabezados))
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, h in enumerate(encabezados):
        hdr[i].text = h
        run = hdr[i].paragraphs[0].runs[0]
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        tc_pr = hdr[i]._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "1A237E")
        tc_pr.append(shd)
    for fila in filas:
        row = t.add_row().cells
        for i, val in enumerate(fila):
            row[i].text = val
    doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# PORTADA
# ══════════════════════════════════════════════════════════════════════════════
titulo_doc("ADAPT-ECG — Notas Técnicas de Sesión")
p = doc.add_paragraph("Sistema inteligente con reentrenamiento continuo para la detección adaptativa de patologías cardiovasculares basado en señales ECG")
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.runs[0].italic = True

doc.add_paragraph()
p2 = doc.add_paragraph("Alumno: Elias Bejarano Lozada (#20213057)\nInstitución: Instituto Tecnológico de Tijuana\nResidencia Profesional — Ingeniería Biomédica\nFecha: 19 de marzo de 2026")
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# TEMA 1 — RED NEURONAL CNN
# ══════════════════════════════════════════════════════════════════════════════
subtitulo("1. Red Neuronal Convolucional (CNN)")

subseccion("¿Por qué se llama CNN?")
parrafo("CNN significa Convolutional Neural Network (Red Neuronal Convolucional). El nombre viene de la operación matemática central que utiliza: la convolución.")
parrafo("Una convolución consiste en deslizar un filtro (kernel) sobre los datos y calcular el producto punto en cada posición. El filtro busca un patrón específico en toda la señal sin importar en qué posición aparezca.")

subseccion("¿Por qué es útil para ECG?")
parrafo("Las señales ECG tienen estructura local: las ondas P, QRS y T son patrones que aparecen en posiciones distintas para cada latido. La convolución:")
bullet("Detecta el mismo patrón sin importar dónde aparezca (invarianza a la traslación)")
bullet("Es eficiente: comparte parámetros en toda la señal")
bullet("Aprende automáticamente qué formas de onda son relevantes")

subseccion("Arquitectura del modelo ECG_CNN")
parrafo("El modelo usa Conv1d porque la entrada es una señal 1D (señal de tiempo). Se define en src/ui/app.py línea 62.")

tabla(
    ["Componente", "Capas", "Qué detecta"],
    [
        ["block1", "Conv1d(1→32, k=5) + BatchNorm + ReLU + MaxPool", "Rasgos básicos de la señal"],
        ["block2", "Conv1d(32→64, k=5) + BatchNorm + ReLU + MaxPool", "Rasgos intermedios"],
        ["block3", "Conv1d(64→128, k=3) + BatchNorm + ReLU + MaxPool", "Rasgos complejos"],
        ["pool", "AdaptiveAvgPool1d(1)", "Resume todo en 1 valor por canal"],
        ["classifier", "Flatten → Linear(128→64) → ReLU → Dropout → Linear(64→5)", "Clasifica en 5 clases AAMI"],
    ]
)

parrafo("La señal entra con forma (batch, 1, 71) y sale como (batch, 5): una probabilidad por clase AAMI.")
parrafo("Total de parámetros entrenables: 44,229.")

subseccion("¿Dónde está el código CNN en el proyecto?")
tabla(
    ["Archivo", "Propósito"],
    [
        ["src/ui/app.py (línea 62)", "Versión en producción — la que corre la UI"],
        ["notebooks/fase2_entrenamiento.ipynb", "Donde se definió y entrenó originalmente"],
        ["notebooks/fase3_reentrenamiento.ipynb", "Donde se realizó el reentrenamiento incremental"],
        ["notebooks/fase4_evaluacion.ipynb", "Donde se evaluó comparativamente"],
    ]
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# TEMA 2 — MODELO BASE VS MODELO REENTRENADO
# ══════════════════════════════════════════════════════════════════════════════
subtitulo("2. Modelo Base vs Modelo Reentrenado")

subseccion("Diferencias generales")
tabla(
    ["Característica", "Modelo Base", "Modelo Reentrenado"],
    [
        ["Archivo", "ecg_cnn_base.pth", "ADAPT-ECG-RETRAINED.pth"],
        ["Fase", "Fase 2", "Fase 3"],
        ["Dónde se entrenó", "Google Colab (GPU T4)", "Google Colab (GPU T4)"],
        ["Cómo se entrenó", "Todos los datos de golpe", "5 lotes secuenciales"],
        ["Épocas", "30 épocas totales", "5 épocas por lote (25 total)"],
        ["Learning rate", "1e-3", "1e-4 (más pequeño, para no olvidar)"],
        ["Técnica especial", "Ninguna", "Replay Buffer (2,000 latidos)"],
        ["Arquitectura CNN", "ECG_CNN (igual)", "ECG_CNN (igual)"],
    ]
)

subseccion("Resultados comparativos")
tabla(
    ["Métrica", "Modelo Base (estático)", "Modelo Reentrenado (adaptativo)"],
    [
        ["Accuracy", "93.54%", "97.83% (+4.29%)"],
        ["F1 Macro", "0.7849", "0.9071 (+0.1222)"],
        ["F1 Weighted", "0.9447", "0.9773 (+0.0326)"],
    ]
)

subseccion("Diferencias en código")
parrafo("Entrenamiento Base (Fase 2 — Google Colab):")
codigo("optimizer = torch.optim.Adam(model.parameters(), lr=1e-3)\ncriterion = nn.NLLLoss(weight=class_weights)  # pesos por desbalance\n\nfor epoch in range(30):\n    for xb, yb in dataloader:\n        optimizer.zero_grad()\n        loss = criterion(model(xb), yb)\n        loss.backward()\n        optimizer.step()")

parrafo("Reentrenamiento Incremental (Fase 3 — función incremental_train, línea 210):")
codigo("def incremental_train(model, beats, labels, replay_buffer, lr=1e-4, epochs=5):\n    optimizer = torch.optim.Adam(model.parameters(), lr=lr)  # lr 10x más pequeño\n    criterion = nn.CrossEntropyLoss()\n    replay_buffer.add(beats, labels)   # guarda en memoria\n\n    for epoch in range(epochs):\n        Xr, yr = replay_buffer.sample(n_replay)   # recuerdos del pasado\n        X_combined = concat(beats, Xr)             # nuevo + recordado\n        # entrena con la mezcla...")

parrafo("La diferencia principal es el Replay Buffer: evita el Catastrophic Forgetting mezclando datos nuevos con muestras de lotes anteriores.")

subseccion("¿Por qué mejora el modelo reentrenado?")
parrafo("El modelo base ve los datos una sola vez. El modelo reentrenado los ve lote por lote y con cada lote:")
bullet("Entrena con datos nuevos del registro actual")
bullet("Mezcla datos nuevos con muestras del Replay Buffer (recuerdos de pacientes anteriores)")
bullet("Actualiza sus pesos sin olvidar lo anterior — resuelve el Catastrophic Forgetting")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# TEMA 3 — MIT-BIH Y AAMI
# ══════════════════════════════════════════════════════════════════════════════
subtitulo("3. Base de Datos MIT-BIH y Estándar AAMI")

subseccion("¿Qué es MIT-BIH?")
parrafo("MIT-BIH Arrhythmia Database — creada por el MIT y el Beth Israel Hospital de Boston. Es la base de datos de referencia estándar para algoritmos de clasificación de arritmias.")

tabla(
    ["Característica", "Valor"],
    [
        ["Número de registros", "48 pacientes reales (numerados 100–234)"],
        ["Duración por registro", "30 minutos de ECG continuo"],
        ["Frecuencia de muestreo", "360 Hz (360 muestras por segundo)"],
        ["Canales por registro", "2 derivaciones de ECG"],
        ["Anotaciones", "Cada latido etiquetado manualmente por dos cardiólogos expertos"],
    ]
)

subseccion("Archivos por registro")
parrafo("Cada registro viene con exactamente 3 archivos (ejemplo con registro 100):")
tabla(
    ["Archivo", "Contenido"],
    [
        ["100.hea", "Cabecera — frecuencia, canales, duración, datos del paciente"],
        ["100.dat", "Señal cruda binaria — voltaje del ECG muestra a muestra"],
        ["100.atr", "Anotaciones — posición y tipo de cada latido (etiquetas del cardiólogo)"],
    ]
)
parrafo("48 registros × 3 archivos = 144 archivos en total en la carpeta local.")

subseccion("Etiquetas originales MIT-BIH")
tabla(
    ["Símbolo", "Significado clínico"],
    [
        ["N, .", "Latido normal"],
        ["L", "Bloqueo de rama izquierda"],
        ["R", "Bloqueo de rama derecha"],
        ["A", "Contracción auricular prematura"],
        ["a", "Contracción auricular prematura aberrante"],
        ["J", "Latido de escape de unión"],
        ["S", "Latido supraventricular prematuro"],
        ["e", "Latido de escape auricular"],
        ["j", "Latido de escape de unión"],
        ["V", "Contracción ventricular prematura (PVC)"],
        ["E", "Latido de escape ventricular"],
        ["F", "Latido de fusión (normal + ventricular)"],
        ["f", "Latido de fusión (marcapasos + normal)"],
        ["/", "Latido de marcapasos"],
        ["!", "Latido ventricular de marcapasos"],
        ["Q", "No clasificable"],
        ["?", "Indefinido"],
    ]
)

subseccion("¿Qué es AAMI?")
parrafo("AAMI = Association for the Advancement of Medical Instrumentation. Es una organización estadounidense que establece estándares para dispositivos médicos. Su estándar EC57 define cómo agrupar los tipos de latidos en 5 categorías clínicamente relevantes.")

subseccion("Conversión MIT-BIH → AAMI (5 clases)")
tabla(
    ["Clase AAMI", "Nombre", "Símbolos MIT-BIH incluidos"],
    [
        ["N", "Normal", "N, ."],
        ["S", "Supraventricular", "A, a, J, S, e, j"],
        ["V", "Ventricular", "V, E"],
        ["F", "Fusión", "F"],
        ["Q", "Desconocido", "Q, ?, /, f, !"],
    ]
)

parrafo("¿Por qué agrupar? Porque clínicamente lo importante es el grupo de arritmia. Además, el estándar AAMI hace los resultados comparables con la literatura científica internacional.")

subseccion("Las etiquetas ya vienen predefinidas")
parrafo("Sí. Cuando se grabaron los ECGs reales, dos cardiólogos expertos revisaron cada latido manualmente y le asignaron un símbolo. Esas anotaciones están en el archivo .atr.")
parrafo("El sistema compara:")
bullet("y_true → etiqueta real del cardiólogo (del archivo .atr convertida a AAMI)")
bullet("y_pred → predicción de la CNN sobre la señal del archivo .dat")
parrafo("La Accuracy que se muestra en pantalla indica en qué porcentaje de latidos la CNN coincidió con el diagnóstico del cardiólogo.")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# TEMA 4 — MATRIZ DE CONFUSIÓN
# ══════════════════════════════════════════════════════════════════════════════
subtitulo("4. Matriz de Confusión")

subseccion("¿Qué es?")
parrafo("Es una tabla que muestra qué tan bien clasificó el modelo, comparando las predicciones contra las etiquetas reales del cardiólogo.")

subseccion("Términos clave")
tabla(
    ["Término", "Significado", "Resultado"],
    [
        ["TP (Verdadero Positivo)", "El modelo dijo V y era V", "✅ Correcto"],
        ["TN (Verdadero Negativo)", "El modelo dijo N y era N", "✅ Correcto"],
        ["FP (Falso Positivo)", "El modelo dijo V pero era N", "❌ Error"],
        ["FN (Falso Negativo)", "El modelo dijo N pero era V", "❌ Error"],
    ]
)

subseccion("Estructura en el proyecto (5 clases AAMI)")
parrafo("La matriz es 5×5 — una fila y columna por cada clase AAMI. Está normalizada por fila, es decir, cada valor representa el porcentaje de latidos reales de esa clase que fueron predichos como cada clase.")
bullet("Diagonal principal (izquierda a derecha) = aciertos → se desea que sean altos (cerca de 1.0)")
bullet("Fuera de la diagonal = errores → se desea que sean bajos (cerca de 0.0)")

subseccion("Métricas derivadas de la matriz")
tabla(
    ["Métrica", "Fórmula", "Qué mide"],
    [
        ["Sensibilidad", "TP / (TP + FN)", "De todos los V reales, ¿cuántos detecté?"],
        ["Especificidad", "TN / (TN + FP)", "De todos los no-V, ¿cuántos classifiqué bien?"],
        ["Precisión", "TP / (TP + FP)", "De los que dije V, ¿cuántos eran realmente V?"],
        ["F1-Score", "2 × Prec × Sens / (Prec + Sens)", "Balance entre precisión y sensibilidad"],
    ]
)

subseccion("¿Por qué es más útil que el Accuracy?")
parrafo("Con clases desbalanceadas (N = 79.3% del dataset) el Accuracy puede ser engañoso. Un modelo que siempre prediga Normal tendría 79% de Accuracy sin aprender nada útil. La matriz muestra exactamente dónde se equivoca el modelo — por ejemplo, si confunde F con V frecuentemente.")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# TEMA 5 — INTERFAZ STREAMLIT
# ══════════════════════════════════════════════════════════════════════════════
subtitulo("5. Interfaz Streamlit — Descripción Completa")

subseccion("Panel lateral (Sidebar) — Configuración")
parrafo("Fuente de datos — radio button con 2 opciones:")
bullet("📂 Base de datos ECG (carpeta) → usa MIT-BIH local")
bullet("📄 CSV propio → sube tu propia señal ECG")

parrafo("Si se elige carpeta:")
bullet("Campo de texto para la ruta (auto-detecta MIT-BIH si existe en el equipo)")
bullet("Mensaje de confirmación con número de registros encontrados")
bullet("Selector desplegable del registro (100, 101, 102...)")
bullet("Checkbox: Filtro pasa-banda (0.5–40 Hz)")

parrafo("Si se elige CSV:")
bullet("Botón para subir archivo .csv")
bullet("Campo para ingresar la frecuencia de muestreo (Hz)")
bullet("Checkbox: Filtro pasa-banda")

parrafo("Selector de modelo:")
bullet("Adaptativo (Reentrenado) — seleccionado por defecto")
bullet("Estático (Base)")

parrafo("Botón ▶ Analizar — color azul, acción principal")
parrafo("Leyenda de clases AAMI con sus colores identificadores.")

subseccion("Pestaña 1 — 🔬 Análisis e Inferencia")
parrafo("Al presionar Analizar aparece:")
bullet("Fila de 5 métricas: Latidos evaluados, Accuracy, F1 Macro, F1 Weighted, Confianza media")
bullet("Gráfica ECG con picos R clasificados por color según clase AAMI")
bullet("Dona con distribución de clases (porcentaje de cada tipo de latido)")
bullet("Matriz de confusión normalizada (heatmap azul)")
bullet("Gráfica de barras agrupadas: Sensibilidad, Especificidad, Precisión, F1 por clase")
bullet("Tabla de métricas por clase con valores exactos")
bullet("Tabla detalle por latido: posición, etiqueta real, predicción, ✅/❌, confianza, probabilidades por clase")
bullet("Botón para descargar resultados en CSV")

parrafo("Sección de Reentrenamiento Continuo:")
bullet("Slider: Épocas (1–15, default 5)")
bullet("Slider: Learning rate (1e-5 hasta 1e-3, default 1e-4)")
bullet("Slider: Tamaño Replay Buffer (500–5000, default 2,000)")
bullet("Métricas ANTES del reentrenamiento para comparar")
bullet("Botón 🔄 Reentrenar modelo adaptativo")
bullet("Al reentrenar: métricas DESPUÉS con delta ±, curva de pérdida por época, distribución ANTES vs DESPUÉS")

subseccion("Pestaña 2 — 📈 Historial de Reentrenamiento")
bullet("Gráfica de línea: Accuracy ANTES vs DESPUÉS por sesión")
bullet("Gráfica de línea: F1 Macro ANTES vs DESPUÉS por sesión")
bullet("Gráfica de barras: Latidos usados por sesión")
bullet("Tabla de sesiones con todas las métricas históricas acumuladas")
bullet("Botón para exportar historial en CSV")
bullet("Botón para borrar el historial")

subseccion("Mapa de código — ¿Dónde está declarado cada elemento?")
tabla(
    ["Línea", "Función / Clase", "Propósito"],
    [
        ["62", "class ECG_CNN", "Define la arquitectura de la red neuronal"],
        ["89", "class ReplayBuffer", "Define la memoria de reentrenamiento incremental"],
        ["122", "load_model_cached()", "Carga modelo con caché de Streamlit"],
        ["128", "load_model_fresh()", "Carga modelo sin caché (post-reentrenamiento)"],
        ["135", "scan_db_folder()", "Detecta registros .hea en la carpeta seleccionada"],
        ["147", "load_record()", "Lee .hea + .dat + .atr con wfdb"],
        ["155", "bandpass_filter()", "Filtro pasa-banda 0.5–40 Hz (Butterworth orden 4)"],
        ["161", "segment_beats_annotated()", "Segmenta 71 muestras por latido con etiquetas"],
        ["181", "segment_beats_csv()", "Segmenta latidos sin etiquetas (modo CSV)"],
        ["199", "run_inference()", "Pasa latidos por la CNN y devuelve predicciones"],
        ["210", "incremental_train()", "Reentrenamiento incremental con Replay Buffer"],
        ["248", "plot_ecg_with_peaks()", "Señal ECG con puntos de colores por clase"],
        ["271", "plot_distribution()", "Dona de distribución de clases"],
        ["282", "plot_confusion()", "Matriz de confusión normalizada"],
        ["299", "plot_per_class_bars()", "Barras de métricas por clase"],
        ["340", "plot_history()", "Gráficas de evolución del historial"],
        ["405", "plot_loss_curve()", "Curva de pérdida del reentrenamiento"],
        ["417", "st.set_page_config()", "Título y layout de la página"],
        ["423", "with st.sidebar:", "Todo el panel izquierdo de configuración"],
        ["490", "st.tabs(...)", "Crea las 2 pestañas principales"],
        ["492", "with tab_history:", "Pestaña 📈 Historial de Reentrenamiento"],
        ["531", "with tab_analysis:", "Pestaña 🔬 Análisis e Inferencia"],
        ["547", "if run_btn:", "Acción al presionar ▶ Analizar"],
    ]
)

subseccion("Flujo completo al presionar Analizar")
parrafo("El sistema ejecuta los siguientes pasos en orden:")
bullet("1. load_record() (línea 550) — lee señal y anotaciones del registro seleccionado")
bullet("2. bandpass_filter() (línea 554) — filtra ruido de la señal ECG")
bullet("3. segment_beats_annotated() (línea 556) — corta 71 muestras por latido usando picos R")
bullet("4. run_inference() (línea 564) — la CNN predice la clase de cada latido")
bullet("5. accuracy_score() (línea 597) — compara predicciones vs etiquetas del cardiólogo")
bullet("6. Funciones plot_* (líneas 612–628) — dibuja todas las gráficas en pantalla")

# ── Pie de página ─────────────────────────────────────────────────────────────
doc.add_page_break()
p = doc.add_paragraph("Documento generado el 19 de marzo de 2026 — ADAPT-ECG · ITT · Ingeniería Biomédica")
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.runs[0].font.color.rgb = RGBColor(0x75, 0x75, 0x75)
p.runs[0].font.size = Pt(9)

# ── Guardar ───────────────────────────────────────────────────────────────────
out = Path("docs/ADAPT-ECG_Notas_Sesion.docx")
doc.save(str(out))
print(f"Documento guardado en: {out}")
