"""
Genera docs/ADAPT-ECG_Informe_Tecnico_CNN.docx
Informe tecnico que explica el flujo y arquitectura CNN del proyecto:
"Sistema inteligente con reentrenamiento continuo para la deteccion adaptativa
de patologias cardiovasculares basado en senales ECG"
Alumno: Elias Bejarano Lozada (#20213057), Ingenieria Biomedica, ITT
Basado en: diagrama_flujo_cnn.drawio y diagrama_cnn_arquitectura.drawio
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = "docs/ADAPT-ECG_Informe_Tecnico_CNN.docx"

# ---------------------------------------------------------------------------
# Helpers de formato (estilo ITT)
# ---------------------------------------------------------------------------

NEGRO = RGBColor(0x00, 0x00, 0x00)

doc = Document()

# Margenes ITT: 3 cm izquierdo, 2.5 cm resto
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(2.5)


def set_run_font(run, size=12, bold=False, italic=False):
    run.font.name   = "Times New Roman"
    run.font.size   = Pt(size)
    run.font.bold   = bold
    run.font.italic = italic
    run.font.color.rgb = NEGRO
    rPr = run._r.get_or_add_rPr()
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), "Times New Roman")
    rFonts.set(qn("w:hAnsi"), "Times New Roman")
    rFonts.set(qn("w:cs"),    "Times New Roman")
    rPr.insert(0, rFonts)


def heading(text, size=13, space_before=14):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.line_spacing = Pt(18)
    run = p.add_run(text)
    set_run_font(run, size=size, bold=True)
    return p


def body(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after       = Pt(6)
    p.paragraph_format.first_line_indent = Cm(0.75)
    p.paragraph_format.line_spacing      = Pt(18)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    set_run_font(run, size=12)
    return p


def bullet(text, indent=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_after  = Pt(3)
    p.paragraph_format.left_indent  = Cm(indent)
    p.paragraph_format.line_spacing = Pt(16)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    set_run_font(run, size=12)
    return p


def tabla_simple(encabezados, filas):
    """Agrega una tabla con encabezados en negrita."""
    table = doc.add_table(rows=1 + len(filas), cols=len(encabezados))
    table.style = "Table Grid"
    # Encabezados
    for i, h in enumerate(encabezados):
        cell = table.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        set_run_font(run, size=11, bold=True)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Filas
    for r, fila in enumerate(filas):
        for c, val in enumerate(fila):
            cell = table.rows[r + 1].cells[c]
            cell.text = ""
            run = cell.paragraphs[0].add_run(val)
            set_run_font(run, size=11)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()


def nota(text):
    """Parrafo de nota tecnica en italicas."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.left_indent  = Cm(0.75)
    p.paragraph_format.line_spacing = Pt(16)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run("Nota: " + text)
    set_run_font(run, size=11, italic=True)
    return p


# ===========================================================================
# PORTADA / ENCABEZADO
# ===========================================================================

p_title = doc.add_paragraph()
p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_title.paragraph_format.space_before = Pt(0)
p_title.paragraph_format.space_after  = Pt(8)
run = p_title.add_run("INFORME TECNICO")
set_run_font(run, size=18, bold=True)

p_sub = doc.add_paragraph()
p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_sub.paragraph_format.space_after = Pt(6)
run = p_sub.add_run(
    "Flujo y Arquitectura de la Red Neuronal Convolucional (CNN)\n"
    "Aplicada a la Deteccion Adaptativa de Arritmias Cardiacas"
)
set_run_font(run, size=14, bold=True)

for line in [
    "Sistema inteligente con reentrenamiento continuo para la deteccion adaptativa",
    "de patologias cardiovasculares basado en senales ECG",
    "",
    "Alumno: Elias Bejarano Lozada  |  No. Control: 20213057",
    "Asesor: M.C. Fortunato Ramirez Arzate",
    "Ingenieria Biomedica — Instituto Tecnologico de Tijuana (ITT)",
    "Residencia Profesional — 2026",
    "",
    "Diagramas de referencia:",
    "diagrama_flujo_cnn.drawio  |  diagrama_cnn_arquitectura.drawio",
]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(line)
    set_run_font(run, size=11, bold=("diagrama" in line.lower() or "Sistema" in line))

doc.add_paragraph()

# ===========================================================================
# 1. INTRODUCCION
# ===========================================================================

heading("1. INTRODUCCION", size=14, space_before=16)

body(
    "El presente informe tecnico describe el flujo de procesamiento y la arquitectura "
    "de la Red Neuronal Convolucional (CNN) disenada para el proyecto ADAPT-ECG. "
    "El sistema tiene como objetivo clasificar latidos cardiacos en cinco categorias "
    "segun el estandar AAMI EC57, a partir de senales ECG provenientes de la base de "
    "datos MIT-BIH Arrhythmia Database. La descripcion se apoya en dos diagramas "
    "elaborados durante el desarrollo del proyecto: el diagrama de flujo general del "
    "sistema (diagrama_flujo_cnn.drawio) y el diagrama detallado de la arquitectura "
    "CNN con su modulo de reentrenamiento continuo (diagrama_cnn_arquitectura.drawio)."
)

body(
    "La contribucion principal del sistema reside en su capacidad de reentrenamiento "
    "incremental: el modelo no es estatico, sino que actualiza sus parametros conforme "
    "llegan nuevos datos de pacientes, sin necesidad de reiniciar el entrenamiento "
    "desde cero y sin incurrir en olvido catastrofico (catastrophic forgetting). "
    "Este informe explica, de forma rigurosa pero accesible, como funciona la "
    "convolucion matematica, que hace cada bloque de la red y como se diferencia "
    "el modelo base del modelo adaptativo."
)

# ===========================================================================
# 2. CONTEXTO: SENALES ECG Y CLASIFICACION AAMI
# ===========================================================================

heading("2. CONTEXTO: SENALES ECG Y CLASIFICACION AAMI", size=13)

body(
    "Una senal ECG registra la actividad electrica del corazon a lo largo del tiempo. "
    "La base de datos MIT-BIH (Massachusetts Institute of Technology — Beth Israel "
    "Hospital) contiene 48 registros de 30 minutos de duracion, muestreados a 360 Hz "
    "y anotados por cardiólogos. Cada anotacion indica la posicion de un pico R "
    "(contraccion ventricular) y el tipo de latido asociado."
)

body(
    "El estandar AAMI EC57 agrupa los tipos de latido en cinco clases que este "
    "sistema debe distinguir:"
)

tabla_simple(
    ["Clase", "Nombre", "Descripcion", "Distribucion en MIT-BIH"],
    [
        ["N", "Normal",            "Latido sinusal normal",                           "79.3 %"],
        ["S", "Supraventricular",  "Origen por encima de los ventriculos",            " 3.2 %"],
        ["V", "Ventricular",       "Contraccion ventricular prematura (PVC)",         " 7.6 %"],
        ["F", "Fusion",            "Mezcla de latido normal y ventricular",           " 0.8 %"],
        ["Q", "Desconocido",       "Marcapasos o no clasificable",                    " 9.0 %"],
    ]
)

body(
    "El sistema procesa en total 94,627 latidos extraidos de los 48 registros. "
    "Cada latido es una ventana de 71 muestras centrada en el pico R: 35 muestras "
    "antes del pico, el pico mismo y 36 muestras despues. Esta ventana equivale a "
    "aproximadamente 197 milisegundos de senal a 360 Hz, tiempo suficiente para "
    "capturar el complejo QRS completo y parte de las ondas P y T adyacentes."
)

# ===========================================================================
# 3. DIAGRAMA 1: FLUJO GENERAL DEL SISTEMA
# ===========================================================================

heading("3. DIAGRAMA 1 — FLUJO GENERAL DEL SISTEMA (diagrama_flujo_cnn.drawio)", size=13)

body(
    "El primer diagrama describe el flujo completo del sistema desde la entrada de "
    "datos hasta la generacion de predicciones en la interfaz Streamlit. El flujo "
    "se organiza en cuatro etapas principales que corresponden a las cuatro fases "
    "del proyecto de residencia."
)

heading("3.1 Entrada de datos", size=12, space_before=10)

body(
    "El sistema acepta dos tipos de entrada. El primero es una carpeta de registros "
    "MIT-BIH en formato nativo de PhysioNet: archivos .hea (cabecera), .dat (senal "
    "digital) y .atr (anotaciones del cardiologo). El segundo es un archivo CSV con "
    "una senal ECG propia del usuario, junto con la frecuencia de muestreo "
    "especificada. En ambos casos, el sistema detecta automaticamente el formato "
    "y aplica el preprocesamiento correspondiente."
)

heading("3.2 Preprocesamiento de la senal", size=12, space_before=10)

body(
    "Antes de ingresar a la CNN, cada senal pasa por tres etapas de preprocesamiento. "
    "Primero, un filtro pasa-banda Butterworth de orden 4 con frecuencias de corte "
    "de 0.5 Hz y 40 Hz elimina el ruido de linea de base (movimiento del electrodo) "
    "y la interferencia de la red electrica (60 Hz). Segundo, se detectan los picos "
    "R mediante la funcion find_peaks de SciPy o mediante las anotaciones oficiales "
    "del archivo .atr. Tercero, alrededor de cada pico R se extrae una ventana de "
    "71 muestras y se normaliza individualmente mediante z-score: se resta la media "
    "de la ventana y se divide entre su desviacion estandar. Esta normalizacion hace "
    "que el modelo sea invariante a la amplitud absoluta de la senal, que varia entre "
    "pacientes y derivaciones."
)

heading("3.3 Inferencia con la CNN", size=12, space_before=10)

body(
    "El conjunto de ventanas normalizadas se agrupa en lotes (batch) y se pasa como "
    "tensor de forma (batch, 1, 71) a la CNN. El modelo procesa todos los latidos "
    "de forma paralela y genera para cada uno un vector de cinco probabilidades. "
    "La clase con mayor probabilidad es la prediccion final. El usuario puede elegir "
    "entre el modelo estatico (ecg_cnn_base.pth) y el modelo adaptativo "
    "(ADAPT-ECG-RETRAINED.pth) desde la interfaz Streamlit."
)

heading("3.4 Reentrenamiento incremental", size=12, space_before=10)

body(
    "Si el usuario activa el modulo de reentrenamiento, el sistema toma los latidos "
    "del registro actual junto con muestras almacenadas en el Replay Buffer y "
    "ejecuta una sesion corta de entrenamiento (5 epocas, lr = 1e-4). El modelo "
    "adaptativo actualiza sus pesos y el buffer se renueva con las nuevas muestras. "
    "Los resultados antes y despues del reentrenamiento se registran en el archivo "
    "retrain_history.json y se visualizan en la pestana de historial de la interfaz."
)

# ===========================================================================
# 4. DIAGRAMA 2: ARQUITECTURA CNN DETALLADA
# ===========================================================================

heading("4. DIAGRAMA 2 — ARQUITECTURA CNN DETALLADA (diagrama_cnn_arquitectura.drawio)", size=13)

body(
    "El segundo diagrama desarrolla en detalle la arquitectura interna de la CNN, "
    "la operacion matematica de la convolucion y el mecanismo de reentrenamiento "
    "continuo. Se organiza en tres secciones: (A) la matematica de la convolucion "
    "1D, (B) el pipeline completo del forward pass y (C) el modulo de reentrenamiento."
)

# ---------------------------------------------------------------------------
# 4.1 Seccion A: La convolucion 1D
# ---------------------------------------------------------------------------

heading("4.1 Seccion A: La Operacion de Convolucion 1D", size=12, space_before=10)

body(
    "La convolucion unidimensional es la operacion fundamental que permite a la CNN "
    "detectar patrones locales en una senal temporal. Para una senal de entrada x[n] "
    "y un filtro (kernel) de pesos w[k] de tamano K, la salida en la posicion n se "
    "calcula como:"
)

p_eq = doc.add_paragraph()
p_eq.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_eq.paragraph_format.space_before = Pt(6)
p_eq.paragraph_format.space_after  = Pt(6)
run = p_eq.add_run("y[n]  =  suma de  w[k] * x[n + k - pad]   para k = 0, 1, ..., K-1")
set_run_font(run, size=12, bold=True)

body(
    "En esta expresion, y[n] es el valor de activacion de salida en la posicion n, "
    "x[n] es la senal ECG de entrada, w[k] son los pesos del filtro que el modelo "
    "aprende durante el entrenamiento mediante retropropagacion, K es el tamano del "
    "kernel (5 en los bloques 1 y 2, 3 en el bloque 3) y pad es el numero de ceros "
    "que se agregan en los extremos de la senal para mantener la longitud de salida "
    "igual a la de entrada (padding = 2 en bloques 1-2, padding = 1 en bloque 3)."
)

body(
    "Para ilustrar el calculo, considere la posicion n = 2 de un latido de ejemplo "
    "con valores x[2] = 0.87, x[3] = 1.25, x[4] = 0.95, x[5] = 0.43, x[6] = -0.12 "
    "y un kernel con pesos w[0] = 0.21, w[1] = -0.45, w[2] = 0.78, w[3] = 0.33, "
    "w[4] = -0.67. El producto punto es:"
)

p_calc = doc.add_paragraph()
p_calc.alignment = WD_ALIGN_PARAGRAPH.LEFT
p_calc.paragraph_format.left_indent  = Cm(1.5)
p_calc.paragraph_format.space_after  = Pt(6)
p_calc.paragraph_format.line_spacing = Pt(16)
run = p_calc.add_run(
    "y[2] = 0.87*(0.21) + 1.25*(-0.45) + 0.95*(0.78) + 0.43*(0.33) + (-0.12)*(-0.67)\n"
    "y[2] = 0.183 - 0.563 + 0.741 + 0.142 + 0.080\n"
    "y[2] = 0.583"
)
set_run_font(run, size=11, bold=False)

body(
    "Este calculo se repite para cada posicion de la senal y para cada uno de los "
    "filtros del bloque. El bloque 1 tiene 32 filtros distintos, por lo que produce "
    "32 mapas de caracteristicas en paralelo. Cada filtro aprende a responder a un "
    "patron diferente de la senal: uno puede especializarse en detectar la pendiente "
    "ascendente del pico R, otro en la repolarizacion de la onda T, otro en la "
    "duracion del segmento QRS, y asi sucesivamente. Esta especializacion emerge "
    "automaticamente durante el entrenamiento, sin que el ingeniero la programe "
    "de forma explicita."
)

heading("4.1.1 BatchNorm, ReLU y MaxPool", size=12, space_before=10)

body(
    "Despues de la convolucion, cada bloque aplica tres operaciones adicionales. "
    "La normalizacion por lote (BatchNorm1d) normaliza las activaciones de cada "
    "canal a media cero y varianza unitaria dentro de cada lote de entrenamiento. "
    "Esto estabiliza el proceso de aprendizaje y permite usar tasas de aprendizaje "
    "mas altas sin que los gradientes exploten o desaparezcan."
)

body(
    "La funcion de activacion ReLU (Rectified Linear Unit) aplica la transformacion "
    "f(x) = max(0, x) elemento a elemento: los valores negativos se convierten en "
    "cero y los positivos se mantienen. Esta operacion introduce no-linealidad en "
    "la red, lo cual es indispensable para que la CNN pueda aprender representaciones "
    "complejas que no son combinaciones lineales simples de la entrada."
)

body(
    "La capa MaxPool1d(2) reduce la longitud de cada mapa de caracteristicas a la "
    "mitad tomando el valor maximo de cada ventana de dos muestras consecutivas. "
    "Esta operacion cumple tres propositos: reduce el numero de parametros y el "
    "costo computacional, agrega invarianza a pequenos desplazamientos del latido "
    "en la ventana y aumenta el campo receptivo efectivo de las capas siguientes."
)

heading("4.1.2 Dropout", size=12, space_before=10)

body(
    "El Dropout(p=0.5) en el clasificador desactiva aleatoriamente el 50% de las "
    "neuronas en cada pasada hacia adelante durante el entrenamiento. Esto obliga "
    "a la red a aprender representaciones redundantes y evita que el modelo "
    "memorice el conjunto de entrenamiento en lugar de generalizar a nuevos datos "
    "(sobreajuste u overfitting). Durante la inferencia, el Dropout se desactiva "
    "automaticamente (modo eval()) y todas las neuronas contribuyen a la prediccion."
)

# ---------------------------------------------------------------------------
# 4.2 Seccion B: Pipeline CNN (Forward Pass)
# ---------------------------------------------------------------------------

heading("4.2 Seccion B: Pipeline Completo del Forward Pass", size=12, space_before=10)

body(
    "El forward pass es el recorrido que realiza un tensor de datos desde la entrada "
    "hasta la salida de la red. La siguiente tabla resume cada etapa del pipeline, "
    "la transformacion que aplica y las dimensiones del tensor resultante. "
    "La notacion (batch, canales, longitud) indica las tres dimensiones del tensor "
    "en cada punto del procesamiento, donde batch es el numero de latidos procesados "
    "en paralelo."
)

tabla_simple(
    ["Etapa", "Operacion", "Tensor de salida", "Parametros"],
    [
        ["Entrada",          "Latido ECG normalizado",                       "(batch, 1, 71)",   "—"],
        ["Block 1 Conv1d",   "Conv1d(1→32, k=5, pad=2)",                     "(batch, 32, 71)",  "192"],
        ["Block 1 BN+ReLU",  "BatchNorm1d(32) + ReLU",                       "(batch, 32, 71)",  "128"],
        ["Block 1 MaxPool",  "MaxPool1d(2) — 71 → 35",                       "(batch, 32, 35)",  "—"],
        ["Block 2 Conv1d",   "Conv1d(32→64, k=5, pad=2)",                    "(batch, 64, 35)",  "10,304"],
        ["Block 2 BN+ReLU",  "BatchNorm1d(64) + ReLU",                       "(batch, 64, 35)",  "256"],
        ["Block 2 MaxPool",  "MaxPool1d(2) — 35 → 17",                       "(batch, 64, 17)",  "—"],
        ["Block 3 Conv1d",   "Conv1d(64→128, k=3, pad=1)",                   "(batch, 128, 17)", "24,704"],
        ["Block 3 BN+ReLU",  "BatchNorm1d(128) + ReLU",                      "(batch, 128, 17)", "512"],
        ["Block 3 MaxPool",  "MaxPool1d(2) — 17 → 8",                        "(batch, 128, 8)",  "—"],
        ["AdaptiveAvgPool",  "AdaptiveAvgPool1d(1) — 8 → 1",                 "(batch, 128, 1)",  "—"],
        ["Flatten",          "Aplana el tensor a vector",                     "(batch, 128)",     "—"],
        ["Linear 1 + ReLU",  "Linear(128→64) + ReLU",                        "(batch, 64)",      "8,256"],
        ["Dropout",          "Dropout(p=0.5)",                                "(batch, 64)",      "—"],
        ["Linear 2",         "Linear(64→5)",                                  "(batch, 5)",       "325"],
        ["Softmax",          "Probabilidades por clase",                      "(batch, 5)",       "—"],
    ]
)

nota(
    "El total de parametros entrenables es 44,229. Este numero relativamente bajo "
    "(comparado con redes como ResNet o VGG) es adecuado para el dominio de ECG "
    "1D con ventanas de 71 muestras. Un modelo mas grande correria el riesgo de "
    "sobreajustarse al dataset de 94,627 latidos."
)

heading("4.2.1 Los tres bloques convolucionales", size=12, space_before=10)

body(
    "La arquitectura organiza las capas convolucionales en tres bloques progresivos. "
    "El primer bloque (Block 1) opera directamente sobre la senal cruda y aprende "
    "caracteristicas de bajo nivel: pendientes, picos y valles locales. Con 32 filtros "
    "de tamano 5, el bloque captura patrones que abarcan aproximadamente 14 ms de "
    "senal a 360 Hz. La salida es un tensor de 32 mapas de caracteristicas de "
    "longitud 35."
)

body(
    "El segundo bloque (Block 2) recibe los 32 mapas del bloque anterior y los "
    "combina mediante 64 filtros de tamano 5. A este nivel, la red aprende "
    "caracteristicas de nivel medio: la forma general del complejo QRS, la "
    "relacion entre la amplitud del pico R y los valles Q y S, y las transiciones "
    "entre diferentes segmentos del latido. La salida es un tensor de 64 mapas "
    "de longitud 17."
)

body(
    "El tercer bloque (Block 3) utiliza 128 filtros de tamano 3 para aprender "
    "representaciones de alto nivel que integran las caracteristicas de los bloques "
    "anteriores. A este nivel, la red distingue entre morfologias globales "
    "correspondientes a las diferentes clases AAMI: un latido ventricular prematuro "
    "tiene una forma caracteristica muy distinta a un latido normal. La salida es "
    "un tensor de 128 mapas de longitud 8."
)

heading("4.2.2 Pooling global y clasificador", size=12, space_before=10)

body(
    "La capa AdaptiveAvgPool1d(1) promedia cada uno de los 128 mapas de "
    "caracteristicas a un unico valor, produciendo un vector de 128 elementos. "
    "Esta operacion elimina la dimension de longitud y hace que la representacion "
    "sea independiente del numero de muestras de la senal original, lo cual "
    "otorga flexibilidad para procesar senales de diferentes duraciones."
)

body(
    "El clasificador (self.classifier) es una red completamente conectada de dos "
    "capas. La primera capa Linear(128, 64) comprime el vector de 128 elementos a "
    "64, seguida de ReLU y Dropout(0.5). La segunda capa Linear(64, 5) proyecta "
    "los 64 valores a 5 logits, uno por clase AAMI. Finalmente, la funcion Softmax "
    "convierte los logits en probabilidades que suman 1.0, lo que permite "
    "interpretar la salida como la confianza del modelo en cada clase."
)

# ---------------------------------------------------------------------------
# 4.3 Seccion C: Reentrenamiento continuo
# ---------------------------------------------------------------------------

heading("4.3 Seccion C: Modulo de Reentrenamiento Continuo", size=12, space_before=10)

body(
    "El modulo de reentrenamiento es la contribucion original del proyecto. "
    "Su objetivo es mantener el rendimiento del modelo conforme llegan nuevos "
    "datos de pacientes, sin reiniciar el entrenamiento desde cero. El diagrama "
    "muestra el flujo completo del proceso, que involucra tres componentes: "
    "el Replay Buffer, la funcion incremental_train() y la logica de "
    "persistencia de pesos."
)

heading("4.3.1 Replay Buffer", size=12, space_before=10)

body(
    "El Replay Buffer es una estructura de datos que almacena hasta 2,000 latidos "
    "de sesiones de entrenamiento anteriores. Su proposito fundamental es evitar "
    "el olvido catastrofico: cuando un modelo se reentrena exclusivamente con datos "
    "nuevos, tiende a sobreescribir los pesos que representaban el conocimiento "
    "anterior, degradando el rendimiento en los datos originales. Al mezclar "
    "muestras antiguas del buffer con los datos nuevos, el modelo sigue siendo "
    "expuesto a la distribucion original y conserva su capacidad de clasificar "
    "correctamente todos los tipos de latidos."
)

body(
    "La implementacion usa una cola circular (FIFO): cuando el buffer supera los "
    "2,000 elementos, los latidos mas antiguos son eliminados para dar lugar a los "
    "mas recientes. El buffer se persiste entre sesiones en el archivo "
    "replay_buffer.npz, de modo que el conocimiento acumulado no se pierde al "
    "cerrar la aplicacion."
)

heading("4.3.2 La funcion incremental_train()", size=12, space_before=10)

body(
    "Cuando el usuario activa el reentrenamiento, la funcion incremental_train() "
    "ejecuta el siguiente procedimiento. Primero, carga el modelo adaptativo actual "
    "(ADAPT-ECG-RETRAINED.pth) sin modificar el modelo base. Segundo, construye "
    "el conjunto de entrenamiento mezclando al 50% los datos nuevos del registro "
    "seleccionado con muestras aleatorias del Replay Buffer, y baraja el resultado "
    "con permutacion aleatoria. Tercero, entrena el modelo durante 5 epocas con "
    "el optimizador Adam a una tasa de aprendizaje de 1e-4, que es diez veces "
    "menor que la usada en el entrenamiento inicial (1e-3). Cuarto, guarda los "
    "nuevos pesos en ADAPT-ECG-RETRAINED.pth y actualiza el buffer."
)

body(
    "La eleccion de lr = 1e-4 es deliberada: una tasa de aprendizaje baja produce "
    "actualizaciones pequenas en los pesos, lo que permite al modelo ajustarse "
    "a los nuevos datos sin destruir las representaciones aprendidas durante el "
    "entrenamiento inicial. Este equilibrio entre plasticidad (capacidad de "
    "aprender cosas nuevas) y estabilidad (capacidad de retener lo aprendido) "
    "es el problema central del aprendizaje continuo (continual learning)."
)

# ===========================================================================
# 5. DIFERENCIA ENTRE MODELO BASE Y MODELO ADAPTATIVO
# ===========================================================================

heading("5. DIFERENCIA ENTRE MODELO BASE Y MODELO ADAPTATIVO", size=13)

body(
    "Una distincion fundamental que el diagrama hace visible es que ambos modelos "
    "comparten exactamente la misma arquitectura ECG_CNN con 44,229 parametros. "
    "La diferencia no reside en la estructura de la red, sino en los valores "
    "numericos de sus pesos. El modelo base fue entrenado una unica vez con todos "
    "los 94,627 latidos durante 30 epocas y sus pesos permanecen fijos. El modelo "
    "adaptativo parte de esos mismos pesos y los refina incrementalmente en sesiones "
    "de 5 epocas cada vez que el usuario introduce un nuevo registro."
)

tabla_simple(
    ["Caracteristica", "Modelo Base", "Modelo Adaptativo"],
    [
        ["Archivo",             "ecg_cnn_base.pth",            "ADAPT-ECG-RETRAINED.pth"],
        ["Arquitectura",        "ECG_CNN (44,229 parametros)",  "ECG_CNN (44,229 parametros)"],
        ["Entrenamiento",       "Una sola vez, 30 epocas",      "Sesion a sesion, 5 epocas c/u"],
        ["Tasa de aprendizaje", "lr = 1e-3",                    "lr = 1e-4 (10x menor)"],
        ["Funcion de perdida",  "Weighted NLLLoss",             "CrossEntropyLoss"],
        ["Se actualiza",        "No",                           "Si, con cada sesion"],
        ["Accuracy (Fase 3)",   "93.54 %",                      "97.83 %"],
        ["F1 Macro (Fase 3)",   "0.7849",                       "0.9071"],
        ["Accuracy (Fase 4)",   "88.72 %",                      "97.41 %"],
        ["F1 Macro (Fase 4)",   "0.7381",                       "0.8954"],
    ]
)

body(
    "Los resultados de la Fase 4 evaluaron ambos modelos sobre 75,702 latidos "
    "correspondientes a los lotes 1 al 4 del dataset. La prueba de McNemar "
    "arrojo un estadistico chi-cuadrado de 5,342 con un valor p aproximadamente "
    "igual a cero, lo que confirma que la diferencia de rendimiento entre ambos "
    "modelos es estadisticamente significativa y no producto del azar."
)

# ===========================================================================
# 6. RESUMEN DEL FLUJO COMPLETO
# ===========================================================================

heading("6. RESUMEN DEL FLUJO COMPLETO", size=13)

body(
    "A continuacion se presenta el flujo completo del sistema en orden cronologico, "
    "desde la adquisicion de datos hasta la prediccion y el reentrenamiento:"
)

pasos = [
    "1. Seleccion del registro MIT-BIH o carga del CSV en la interfaz Streamlit.",
    "2. Lectura de la senal con wfdb.rdrecord() y las anotaciones con wfdb.rdann().",
    "3. Filtrado pasa-banda Butterworth (0.5 Hz - 40 Hz) para eliminar ruido.",
    "4. Deteccion de picos R mediante anotaciones .atr o busqueda automatica con find_peaks.",
    "5. Segmentacion: extraccion de ventanas de 71 muestras centradas en cada pico R.",
    "6. Normalizacion z-score individual por latido: (beat - media) / desviacion estandar.",
    "7. Apilamiento en tensor (batch, 1, 71) para procesamiento en paralelo.",
    "8. Forward pass por ECG_CNN: Block1 → Block2 → Block3 → AvgPool → Flatten → Classifier.",
    "9. Aplicacion de Softmax: conversion de logits a probabilidades (suma = 1.0).",
    "10. Seleccion de la clase con mayor probabilidad como prediccion final.",
    "11. (Opcional) Reentrenamiento: mezcla buffer + nuevos datos → 5 epocas → actualiza modelo.",
    "12. Registro de metricas en retrain_history.json y visualizacion en la interfaz.",
]

for paso in pasos:
    bullet(paso, indent=1.2)

doc.add_paragraph()

# ===========================================================================
# 7. NOTAS TECNICAS
# ===========================================================================

heading("7. NOTAS TECNICAS IMPORTANTES", size=13)

nota(
    "BEAT_LEN = 71: siempre usar este valor. El archivo settings.py menciona 72 "
    "en algunos contextos, pero el valor operativo correcto es 71 "
    "(35 antes del pico R + el pico R mismo + 36 despues)."
)

nota(
    "La capa clasificadora se llama self.classifier en el codigo (no self.fc). "
    "Las claves de los pesos en el state_dict son classifier.1 para Linear(128,64) "
    "y classifier.4 para Linear(64,5)."
)

nota(
    "El registro 124 de MIT-BIH tiene solo 88 latidos (paciente con marcapasos). "
    "No es adecuado para demostrar reentrenamiento. Para demos se recomienda "
    "usar los registros 100, 200 o 208, que tienen mas de 2,000 latidos y "
    "variedad de clases."
)

nota(
    "La interfaz Streamlit se inicia con el comando: "
    ".venv\\Scripts\\streamlit run src\\ui\\app.py"
    " desde la carpeta raiz del proyecto."
)

# ===========================================================================
# 8. CONCLUSION
# ===========================================================================

heading("8. CONCLUSION", size=13)

body(
    "El sistema ADAPT-ECG demuestra que una CNN de tamano moderado (44,229 "
    "parametros) es suficiente para clasificar arritmias cardiacas con alta "
    "precision cuando se entrena correctamente sobre el dataset MIT-BIH. "
    "La arquitectura de tres bloques convolucionales progresivos, combinada con "
    "BatchNorm, ReLU y MaxPool, permite extraer jerarquicamente las caracteristicas "
    "relevantes de la senal ECG: desde patrones locales de bajo nivel en el "
    "primer bloque hasta representaciones morfologicas globales en el tercero."
)

body(
    "La contribucion principal, el modulo de reentrenamiento continuo con Replay "
    "Buffer, resuelve el problema del concept drift en datos medicos: los modelos "
    "estaticos degradan su rendimiento cuando los nuevos datos difieren de la "
    "distribucion original, pero el modelo adaptativo mantiene e incluso mejora "
    "su rendimiento gracias a la actualizacion incremental de parametros. "
    "La mejora de 8.69 puntos porcentuales en accuracy y 0.1573 en F1 Macro "
    "sobre 75,702 latidos, con significancia estadistica confirmada por la prueba "
    "de McNemar, valida la efectividad del enfoque propuesto."
)

body(
    "Los diagramas diagrama_flujo_cnn.drawio y diagrama_cnn_arquitectura.drawio "
    "constituyen una representacion visual completa del sistema que sintetiza tanto "
    "el flujo de datos como la matematica subyacente, y sirven como material de "
    "apoyo para la presentacion y defensa del proyecto de residencia profesional."
)

# ===========================================================================
# REFERENCIAS
# ===========================================================================

heading("REFERENCIAS", size=13)

refs = [
    "[1] Moody, G. B. y Mark, R. G. (2001). The impact of the MIT-BIH Arrhythmia "
    "Database. IEEE Engineering in Medicine and Biology Magazine, 20(3), 45-50.",

    "[2] Association for the Advancement of Medical Instrumentation. (2012). "
    "ANSI/AAMI EC57: Testing and reporting performance results of cardiac rhythm "
    "and ST-segment measurement algorithms. Arlington, VA: AAMI.",

    "[3] LeCun, Y., Bengio, Y. y Hinton, G. (2015). Deep learning. Nature, "
    "521(7553), 436-444.",

    "[4] Ioffe, S. y Szegedy, C. (2015). Batch normalization: Accelerating deep "
    "network training by reducing internal covariate shift. Proceedings of the "
    "32nd International Conference on Machine Learning (ICML), 448-456.",

    "[5] Srivastava, N., Hinton, G., Krizhevsky, A., Sutskever, I. y Salakhutdinov, "
    "R. (2014). Dropout: A simple way to prevent neural networks from overfitting. "
    "Journal of Machine Learning Research, 15(1), 1929-1958.",

    "[6] Rolnick, D., Ahuja, A., Schwarz, J., Lillicrap, T. y Wayne, G. (2019). "
    "Experience replay for continual learning. Advances in Neural Information "
    "Processing Systems (NeurIPS), 32.",

    "[7] Paszke, A., Gross, S., Massa, F., et al. (2019). PyTorch: An imperative "
    "style, high-performance deep learning library. Advances in Neural Information "
    "Processing Systems (NeurIPS), 32.",

    "[8] Goldberger, A. L., Amaral, L. A., Glass, L., et al. (2000). PhysioBank, "
    "PhysioToolkit, and PhysioNet: Components of a new research resource for "
    "complex physiologic signals. Circulation, 101(23), e215-e220.",
]

for r in refs:
    p = doc.add_paragraph()
    p.paragraph_format.space_after       = Pt(4)
    p.paragraph_format.left_indent       = Cm(1.0)
    p.paragraph_format.first_line_indent = Cm(-1.0)
    p.paragraph_format.line_spacing      = Pt(15)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(r)
    set_run_font(run, size=11)

# ===========================================================================
# GUARDAR
# ===========================================================================

doc.save(OUT)
print(f"Documento generado: {OUT}")
