"""
Genera el informe de Residencia Profesional ADAPT-ECG como documento Word (.docx)
combinando Resumen, Cap1, Cap2, Cap3 y Cap4.
"""
import re, os, sys
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE = r"C:\Users\Marti\Desktop\PROYECT_ADAPT-ECG\Residencia_Profesional (1)"
OUT  = os.path.join(BASE, "ADAPT-ECG_Residencia_Profesional.docx")

# ─────────────────────────────────────────────
# LIMPIEZA LATEX → TEXTO PLANO
# ─────────────────────────────────────────────
def clean(text):
    # metadatos TCIDATA
    text = re.sub(r'^%TCIDATA.*$', '', text, flags=re.MULTILINE)
    # comentarios
    text = re.sub(r'(?<!\\)%.*$', '', text, flags=re.MULTILINE)
    # tolerancias
    text = re.sub(r'\\(pre)?tolerance=\d+', '', text)
    # \label{} → vacío
    text = re.sub(r'\\label\{[^}]+\}', '', text)
    # \ref{} → (ref)
    text = re.sub(r'\\ref\{[^}]+\}', '(ref)', text)
    # \cite{key} → [key]
    text = re.sub(r'\\cite\{([^}]+)\}',
                  lambda m: '[' + m.group(1).split(',')[0].strip() + ']', text)
    # entornos figure → [FIGURA POR INSERTAR]
    text = re.sub(r'\\begin\{figure\}.*?\\end\{figure\}',
                  '\n[FIGURA POR INSERTAR]\n', text, flags=re.DOTALL)
    # entornos table / tabular → simplificar
    text = re.sub(r'\\begin\{table\}[^\n]*\n?', '\n', text)
    text = re.sub(r'\\end\{table\}', '\n', text)
    text = re.sub(r'\\begin\{tabular\}\{[^}]+\}', '', text)
    text = re.sub(r'\\end\{tabular\}', '', text)
    text = re.sub(r'\\hline', '', text)
    text = re.sub(r'\\multicolumn\{[^}]+\}\{[^}]+\}\{([^}]+)\}', r'\1', text)
    # ecuaciones
    text = re.sub(r'\\begin\{equation\}(.*?)\\end\{equation\}', r'\1', text, flags=re.DOTALL)
    text = re.sub(r'\\begin\{(align|gather|eqnarray)\*?\}(.*?)\\end\{(align|gather|eqnarray)\*?\}',
                  r'\2', text, flags=re.DOTALL)
    text = re.sub(r'\\(nonumber|label)\{[^}]*\}', '', text)
    # math inline $...$ → contenido
    text = re.sub(r'\$\\chi\^2_?\{?[^}]*\}?\$', 'χ²', text)
    text = re.sub(r'\$p\\s*\\\\approx\s*0\$', 'p ≈ 0', text)
    text = re.sub(r'\$([^$\n]{1,120})\$', r'\1', text)
    # entornos simples
    for env in ['center','flushright','flushleft','enumerate','itemize']:
        text = re.sub(r'\\begin\{'+env+r'\}', '', text)
        text = re.sub(r'\\end\{'+env+r'\}', '', text)
    # \item
    text = re.sub(r'\\item\s*', '\n• ', text)
    # comandos con argumento → conservar argumento
    for cmd in ['textbf','textit','texttt','emph','textrm','textsf',
                'textsc','underline','caption','footnote','nonumchapter']:
        text = re.sub(r'\\'+cmd+r'\{([^{}]*)\}', r'\1', text)
    # \href{url}{texto}
    text = re.sub(r'\\href\{[^}]+\}\{([^}]+)\}', r'\1', text)
    # \url{...}
    text = re.sub(r'\\url\{[^}]+\}', '', text)
    # comandos sin argumento
    text = re.sub(r'\\(bigskip|medskip|smallskip|noindent|clearpage|newpage|appendix)\b\s*', '\n', text)
    text = re.sub(r'\\(setlinespacing|decimalpoint|deactivatetilden)\{[^}]*\}', '', text)
    text = re.sub(r'\\(setlinespacing|decimalpoint|deactivatetilden)\b\s*', '', text)
    text = re.sub(r'\\rule\{[^}]+\}\{[^}]+\}', '─────────────────────────────', text)
    # \\ → salto de línea (en tablas)
    text = text.replace('\\\\', ' | ')
    # & → separador columna
    text = text.replace(' & ', ' | ')
    # comandos residuales \cmd{arg} multinivel → arg
    for _ in range(4):
        text = re.sub(r'\\[a-zA-Z*]+\{([^{}]*)\}', r'\1', text)
    # comandos sin argumento residuales
    text = re.sub(r'\\[a-zA-Z*]+\b\s*', '', text)
    # llaves sueltas
    text = text.replace('{','').replace('}','')
    # limpiar líneas vacías múltiples
    text = re.sub(r'\n{3,}', '\n\n', text)
    return text.strip()


# ─────────────────────────────────────────────
# CONFIGURACIÓN DEL DOCUMENTO
# ─────────────────────────────────────────────
def setup_doc():
    doc = Document()
    for sec in doc.sections:
        sec.top_margin    = Cm(3.0)
        sec.bottom_margin = Cm(2.5)
        sec.left_margin   = Cm(3.0)
        sec.right_margin  = Cm(2.5)

    # Estilo Normal → Times New Roman 12pt, interlineado 1.5
    ns = doc.styles['Normal']
    ns.font.name = 'Times New Roman'
    ns.font.size = Pt(12)
    ns.paragraph_format.line_spacing    = Pt(18)
    ns.paragraph_format.space_after     = Pt(0)
    ns.paragraph_format.space_before    = Pt(0)

    # Heading 1 → Capítulo
    h1 = doc.styles['Heading 1']
    h1.font.name = 'Times New Roman'
    h1.font.size = Pt(14)
    h1.font.bold = True
    h1.font.color.rgb = RGBColor(0,0,0)
    h1.paragraph_format.space_before = Pt(24)
    h1.paragraph_format.space_after  = Pt(12)

    # Heading 2 → Sección
    h2 = doc.styles['Heading 2']
    h2.font.name = 'Times New Roman'
    h2.font.size = Pt(13)
    h2.font.bold = True
    h2.font.color.rgb = RGBColor(0,0,0)
    h2.paragraph_format.space_before = Pt(18)
    h2.paragraph_format.space_after  = Pt(6)

    # Heading 3 → Subsección
    h3 = doc.styles['Heading 3']
    h3.font.name = 'Times New Roman'
    h3.font.size = Pt(12)
    h3.font.bold = True
    h3.font.italic = True
    h3.font.color.rgb = RGBColor(0,0,0)
    h3.paragraph_format.space_before = Pt(12)
    h3.paragraph_format.space_after  = Pt(6)

    return doc


# ─────────────────────────────────────────────
# PORTADA
# ─────────────────────────────────────────────
def add_portada(doc):
    def centro(texto, bold=False, size=12):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(texto)
        run.bold = bold
        run.font.name = 'Times New Roman'
        run.font.size = Pt(size)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = Pt(18)
        return p

    centro('SEP                         TECNOLÓGICO NACIONAL DE MÉXICO', bold=True, size=13)
    centro('INSTITUTO TECNOLÓGICO DE TIJUANA', bold=True, size=13)
    centro('DEPARTAMENTO DE INGENIERÍA ELÉCTRICA Y ELECTRÓNICA', size=12)
    doc.add_paragraph()
    doc.add_paragraph()
    centro('RESIDENCIA PROFESIONAL', size=13)
    doc.add_paragraph()
    centro('SISTEMA INTELIGENTE CON REENTRENAMIENTO CONTINUO PARA LA', bold=True, size=13)
    centro('DETECCIÓN ADAPTATIVA DE PATOLOGÍAS CARDIOVASCULARES', bold=True, size=13)
    centro('BASADO EN SEÑALES ECG', bold=True, size=13)
    doc.add_paragraph()
    doc.add_paragraph()

    def derecha(texto, bold=False):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run(texto)
        run.bold = bold
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = Pt(18)

    derecha('PRESENTA:')
    derecha('ELIAS BEJARANO LOZADA', bold=True)
    derecha('20213057', bold=True)
    doc.add_paragraph()
    derecha('BAJO LA ASESORÍA:')
    derecha('INTERNA: M.C. [NOMBRE ASESOR INTERNO]')
    derecha('EXTERNA: M.C. FORTUNATO RAMÍREZ ARZATE')
    doc.add_paragraph()
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run('MARZO DEL 2026                                          TIJUANA, B.C., MÉXICO')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

    doc.add_page_break()


# ─────────────────────────────────────────────
# PÁRRAFO DE TEXTO FORMATEADO
# ─────────────────────────────────────────────
def add_paragraph(doc, text, bold_markers=True):
    """Agrega un párrafo con soporte básico de negritas (**...**) e itálicas (*...*) """
    if not text.strip():
        return
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = Pt(18)
    p.paragraph_format.space_after  = Pt(0)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Dividir por marcadores de negrita/itálica residuales
    # Por simplicidad, añadir como run único
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.line_spacing = Pt(18)
    p.paragraph_format.space_after  = Pt(0)
    run = p.add_run(text.lstrip('• ').strip())
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)


def add_numbered(doc, text, n):
    p = doc.add_paragraph(style='List Number')
    p.paragraph_format.line_spacing = Pt(18)
    p.paragraph_format.space_after  = Pt(0)
    run = p.add_run(text.strip())
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)


# ─────────────────────────────────────────────
# PARSER PRINCIPAL: TEX → WORD
# ─────────────────────────────────────────────
def tex_to_doc(doc, filepath, chapter_num=None):
    with open(filepath, encoding='utf-8', errors='replace') as f:
        raw = f.read()

    text = clean(raw)
    lines = text.split('\n')

    in_enum = False
    enum_counter = 0
    in_bullet = False

    for line in lines:
        line = line.rstrip()

        # Saltar líneas vacías o solo espacios
        if not line.strip():
            in_enum = False
            in_bullet = False
            enum_counter = 0
            continue

        # ── Capítulo ──────────────────────────────
        m = re.match(r'\\chapter\{(.+)\}', line)
        if m:
            title = m.group(1).strip()
            if chapter_num:
                title = f'Capítulo {chapter_num}. {title}'
            doc.add_heading(title, level=1)
            continue

        # Capítulo sin número (\nonumchapter ya fue limpiado, queda el texto)
        # capítulos no numerados como "Resumen" o "Agradecimientos"
        # → manejado abajo como texto libre, pero si la línea dice solo el título:

        # ── Sección ───────────────────────────────
        m = re.match(r'\\section\{(.+)\}', line)
        if m:
            doc.add_heading(m.group(1).strip(), level=2)
            continue

        # ── Subsección ────────────────────────────
        m = re.match(r'\\subsection\{(.+)\}', line)
        if m:
            doc.add_heading(m.group(1).strip(), level=3)
            continue

        # ── Bullet ────────────────────────────────
        if line.strip().startswith('•'):
            add_bullet(doc, line.strip())
            continue

        # ── Tabla simplificada (líneas con |) ─────
        if ' | ' in line and line.strip():
            p = doc.add_paragraph()
            p.paragraph_format.line_spacing = Pt(18)
            run = p.add_run(line.strip())
            run.font.name = 'Courier New'
            run.font.size = Pt(10)
            continue

        # ── [FIGURA POR INSERTAR] ─────────────────
        if '[FIGURA POR INSERTAR]' in line:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run('[FIGURA POR INSERTAR]')
            run.font.name = 'Times New Roman'
            run.font.size = Pt(11)
            run.font.italic = True
            run.font.color.rgb = RGBColor(150, 150, 150)
            continue

        # ── Separador ─────────────────────────────
        if '─────' in line:
            p = doc.add_paragraph()
            run = p.add_run('─' * 60)
            run.font.size = Pt(8)
            continue

        # ── Texto normal ──────────────────────────
        if line.strip():
            add_paragraph(doc, line.strip())


# ─────────────────────────────────────────────
# SECCIÓN ESPECIAL: RESUMEN
# ─────────────────────────────────────────────
def add_resumen(doc):
    doc.add_heading('Resumen', level=1)
    tex_to_doc(doc, os.path.join(BASE, 'Resumen.tex'))
    doc.add_page_break()


# ─────────────────────────────────────────────
# MAIN
# ─────────────────────────────────────────────
def main():
    print("Generando documento Word...")
    doc = setup_doc()

    # Portada
    add_portada(doc)

    # Resumen
    add_resumen(doc)

    # Capítulos
    chapters = [
        ('Cap1_Introduccion.tex',  1),
        ('Cap2_Marco_Teorico.tex', 2),
        ('Cap3_Resultados.tex',    3),
        ('Cap4_Conclusiones.tex',  4),
    ]

    for filename, num in chapters:
        path = os.path.join(BASE, filename)
        print(f"  Procesando {filename}...")
        tex_to_doc(doc, path, chapter_num=num)
        doc.add_page_break()

    # Guardar
    doc.save(OUT)
    print(f"\nDocumento guardado en:\n  {OUT}")

if __name__ == '__main__':
    main()
